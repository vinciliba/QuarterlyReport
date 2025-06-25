import os
import sys
import sqlite3, json, textwrap, pathlib, base64
from jinja2 import Environment, FileSystemLoader, select_autoescape, DebugUndefined
import streamlit as st
from streamlit.components.v1 import html    
import pandas as pd
from datetime import datetime
import traceback # Import traceback for better error reporting
from ingestion.db_utils import (
    get_all_reports, create_new_report,
    define_expected_table, get_suggested_structure,
    load_report_params, upsert_report_param, save_report_object,
    get_report_object, list_report_objects, delete_report_object, get_variable_status,
    fetch_vars_for_report, compute_cutoff_related_dates, fetch_gt_image, insert_variable, get_existing_rule_for_report
)
from ingestion.report_check import check_report_readiness
import io, docx
import pyperclip
from pathlib import Path
from io import BytesIO
from ui_helpers.ui_tables_helpers import *


DB_PATH = 'database/reporting.db'


# Adjust path as needed
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# =================================
# === Session State Initialization ===
# =================================
# Initialize session state variables used across reruns

# For multi-select report metadata deletion confirmation
if 'pending_delete_reports' not in st.session_state:
    st.session_state.pending_delete_reports = None

# For multi-select upload record deletion confirmation
if 'pending_delete_uploads' not in st.session_state:
    st.session_state.pending_delete_uploads = None

# For auto-selecting the report after a new one is created in Single Upload tab
if 'selected_report_after_create' not in st.session_state:
    st.session_state.selected_report_after_create = None

# Add this new state variable to track the file uploader state
if 'file_uploader_key_counter' not in st.session_state:
    st.session_state.file_uploader_key_counter = 0

# Track completed modules for staged report generation
if 'completed_modules' not in st.session_state:
    st.session_state.completed_modules = []

# Track the last chosen report to persist across reruns
if 'last_chosen_report' not in st.session_state:
    st.session_state.last_chosen_report = None

# Track the last cutoff date to persist across reruns
if 'last_cutoff_date' not in st.session_state:
    st.session_state.last_cutoff_date = None
# =================================
# === End Session State Init ===
# =================================

# --- Rest of your script starts here ---
# DB_PATH = 'database/reporting.db'
# os.makedirs("app_files", exist_ok=True)
# init_db(db_path=DB_PATH)
# etc.
# Use absolute imports for clarity and robustness
try:
    from ingestion.db_utils import (
        init_db,
        get_existing_rule,
        insert_sheet_rule,
        get_transform_rules,
        save_transform_rules,
        insert_upload_log,
        create_new_report,
        get_all_reports,
        is_report_complete,
        get_expected_tables,
        # Import other necessary functions from db_utils
        register_file_alias, # Added based on likely use
        get_alias_for_file, # Added
        update_alias_status, # Added
        get_alias_last_load, # Added
        log_cutoff # Added
    )
except ImportError as e:
    st.error(f"Failed to import db_utils: {e}")
    st.stop() # Stop execution if core imports fail


os.makedirs("app_files", exist_ok=True)

# --- Page config, style ---
st.set_page_config(layout="centered")
st.markdown("""
<style>
.main .block-container {
    max-width: 1000px;
    margin: auto;
}
.stToast {
    max-width: 500px;
    margin: 0 auto;
}
</style>
""", unsafe_allow_html=True)

# --- Init DB ---
init_db(db_path=DB_PATH)

# --- Navigation Selectbox ---
# Define your sections
sections = {
    "🚀 Choose Workflow": "workflow",
    "📤 Export Report"  : "export_report",
    "📂 Single File Upload": "single_upload",
    "📦 Mass Upload": "mass_upload",
    "🔎 View History": "history",
    "🛠️ Admin" :  "report_structure",
    "🖋️ Template Editor": "template_editor",
    "📊 Audit Data Input": "audit_data_input"  # New section
    
}

# Use a selectbox for navigation
st.sidebar.title("Navigation")
selected_section_key = st.sidebar.selectbox("Go to", list(sections.keys()))
selected_section = sections[selected_section_key]

# --- Content Area based on Selection ---

# Add a debug print to confirm script execution and selected section
print(f"DEBUG: Script rerun. Selected section: {selected_section_key} ({selected_section})")
st.write(f"", unsafe_allow_html=True) # HTML comment for browser source check

###############################################################################
# Helper – pretty‑print raw variable values                                #####
###############################################################################

def _pretty_print_value(value: str | None) -> None:
    """Render *value* nicely in Streamlit (JSON, table, text)."""
    if value is None:
        st.info("No stored value.")
        return

    try:
        parsed = json.loads(value)
    except (TypeError, ValueError):
        parsed = value

    if isinstance(parsed, list) and parsed and isinstance(parsed[0], dict):
        st.dataframe(pd.DataFrame(parsed), use_container_width=True)
    elif isinstance(parsed, dict):
        st.json(parsed, expanded=False)
    else:
        st.code(str(parsed))

# Helper function to fetch raw value
def _fetch_raw_value(report: str, var: str) -> str | None:
    q = "SELECT value FROM report_variables WHERE report_name = ? AND var_name = ? ORDER BY created_at DESC LIMIT 1"
    with sqlite3.connect(DB_PATH) as con:
        row = con.execute(q, (report, var)).fetchone()
    return row[0] if row else None

# ──────────────────────────────────────────────────
# WORKFLOW – Launch & Validation (Refactored)
# ──────────────────────────────────────────────────
if selected_section == "workflow":
    from ingestion.report_check import check_report_readiness
    import importlib
    from pathlib import Path
    from docx import Document
    from ingestion.db_utils import list_report_modules

    st.title("📊 Report Launch & Validation")

    # Step 1: Pick a report
    reports_df = get_all_reports(DB_PATH)
    if reports_df.empty:
        st.info("No reports defined yet. Create one in “Single File Upload”.")
        st.stop()

    default_report = st.session_state.last_chosen_report if st.session_state.last_chosen_report in reports_df["report_name"].tolist() else reports_df["report_name"].iloc[0]
    chosen_report = st.selectbox("Choose report to validate", reports_df["report_name"].tolist(), index=reports_df["report_name"].tolist().index(default_report))
    st.session_state.last_chosen_report = chosen_report

    # Step 2: Choose cutoff & tolerance
    default_cutoff = st.session_state.last_cutoff_date if st.session_state.last_cutoff_date else datetime.today()
    cutoff_date = st.date_input("📅 Reporting cutoff date", value=default_cutoff)
    st.session_state.last_cutoff_date = cutoff_date
    tolerance_days = st.slider("⏱️ Tolerance (days before cutoff)", 0, 15, 3)

    # Step 3: Validate readiness
    validation_df, ready = check_report_readiness(chosen_report, cutoff_date, tolerance_days, db_path=DB_PATH)
    st.markdown("### Validation results")
    st.dataframe(validation_df, hide_index=True, use_container_width=True)

    if not ready:
        st.error("⛔ Missing or stale uploads detected. Fix them before launch.")
        st.markdown("Please go to the **Single File Upload** or **Mass Upload** section to upload the required data files.")
        st.stop()

    st.success("🎉 All required tables uploaded and within the valid window!")
    
    # Step 3.1: Compute date-based report parameters
    st.markdown("### 📅 Derived Reporting Dates")
    if st.button("📆 Generate Derived Dates"):
        from ingestion.db_utils import compute_cutoff_related_dates, upsert_report_param
        derived = compute_cutoff_related_dates(cutoff_date)
        for k, v in derived.items():
            upsert_report_param(chosen_report, k, v, DB_PATH)
        insert_variable(
            report=chosen_report,
            module="DateParams",
            var="quarter_period",
            value=derived["quarter_period"],
            db_path=DB_PATH,
            anchor="quarter_period"
        )
        st.success("✅ Derived date-based parameters saved to report_params.")

    # Step 3.2: Amendment Report Date hardcoding 
    st.subheader("📅 Select Amendments Report Date")
    amd_report_date = st.date_input("Choose date:", value=datetime.today())
    if st.button("💾 Save Amendments Date"):
        upsert_report_param(chosen_report, "amendments_report_date", amd_report_date.isoformat(), DB_PATH)
        st.success("✅ Amendments date saved.")

    # Step 4: Show available DOCX template
    template_dir = Path("reporting/templates/docx")
    template_file = next((f for f in template_dir.glob("*.docx") if chosen_report.replace(" ", "_") in f.name), None)
    if template_file:
        st.markdown(f"🖋️ **Template in use:** `{template_file.name}`")
    else:
        st.warning(f"No template found for `{chosen_report}` in `/templates/docx/`")
        st.stop()

    # Step 5: Load report module + registry and fetch saved modules
    report_to_module = {
        "Quarterly_Report": "reporting.quarterly_report",
        "Invoice_Summary": "reporting.invoice_summary",
    }

    mod_path = report_to_module.get(chosen_report)
    if not mod_path:
        st.warning(f"No Python module path mapped for `{chosen_report}`.")
        st.stop()

    try:
        mod = importlib.import_module(mod_path)
        print(f"Loaded module: {mod}, dir: {dir(mod)}")  # Debug print
        MODULES = getattr(mod, "MODULES", {})
    except Exception as e:
        st.error(f"Error loading report modules: {e}")
        st.stop()

    if not MODULES:
        st.warning("No modules found for this report.")
        st.stop()

    # Fetch modules from the report_modules table
    saved_modules_df = list_report_modules(chosen_report, DB_PATH)
    saved_module_names = saved_modules_df['module_name'].tolist() if not saved_modules_df.empty else []

    # Step 6: Choose subset of modules
    st.markdown("### 🧩 Select modules to run")
    if not saved_module_names:
        st.warning("Please add modules for the report running in the 'Manage Report-Module Mappings' section.")
        run_button_visible = False
    else:
        remaining_modules = [k for k in saved_module_names if k not in st.session_state.completed_modules]
        enabled_module_names = st.multiselect(
            "Choose modules",
            saved_module_names,
            default=remaining_modules,
            key=f"multiselect_{chosen_report}"
        )
        selected_modules = {k: MODULES.get(k) for k in enabled_module_names}
        run_button_visible = True

    # Step 7: Show progress of completed modules
    st.markdown("### 📈 Progress")
    if st.session_state.completed_modules:
        st.write(f"Completed modules: {', '.join(st.session_state.completed_modules)}")
    else:
        st.write("No modules completed yet.")

    # Step 8: Option to reset the staged report
    if st.button("🧹 Reset staged report"):
        if template_file:
            st.session_state.staged_docx = Document(str(template_file))
            st.toast("Staged DOCX initialized from template.", icon="📝")
        else:
            st.session_state.staged_docx = Document()
            st.toast("Staged DOCX initialized empty.", icon="🆕")
        st.session_state.completed_modules = []
        st.toast("Progress reset.", icon="🔄")

    # Step 9: Run report with selected modules
    if run_button_visible and st.button("🚀 Run Report"):
        if not selected_modules:
            st.warning("Please select at least one module to run.")
        else:
            status = st.empty()
            status.info(f"Launching report **{chosen_report}** …")

            try:
                with st.spinner("Running selected modules…"):
                    if not hasattr(mod, "run_report"):
                        raise RuntimeError(f"Module `{mod_path}` has no `run_report()`")
                    ctx, run_results = mod.run_report(
                        cutoff_date=cutoff_date,
                        tolerance=tolerance_days,
                        db_path=DB_PATH,
                        selected_modules=selected_modules
                    )
            except Exception as e:
                status.empty()
                st.error(f"💥 Error running report: {e}")
                st.code(traceback.format_exc())
            else:
                status.empty()
                st.markdown("### 📝 Module Run Results")
                all_ok = True
                for mod_name, state, msg in run_results:
                    if state == "✅ Success":
                        st.success(f"{mod_name}: {state}")
                        if mod_name not in st.session_state.completed_modules:
                            st.session_state.completed_modules.append(mod_name)
                    else:
                        st.error(f"{mod_name}: {state}")
                        if msg:
                            st.code(msg)
                        all_ok = False
                if all_ok:
                    st.toast("Report finished", icon="✅")
                    st.success(f"Report **{chosen_report}** completed successfully.")
                    if hasattr(mod, "render_and_export"):
                        final_report_path = mod.render_and_export(chosen_report, cutoff_date, ctx)
                        st.success(f"Final report saved as `{final_report_path}` in `app_files/`")
                    else:
                        st.warning(f"No render_and_export function found for `{chosen_report}`.")
                else:
                    st.warning("Some modules failed. See details above.")
                log_cutoff(
                    chosen_report,
                    f"Validation_{cutoff_date}",
                    cutoff_date.isoformat(),
                    validated=True,
                    db_path=DB_PATH,
                )

###############################################################################
# EXPORT REPORT TAB                                                       #####
###############################################################################
if selected_section == "export_report":
    from PIL import Image
    from docxtpl import DocxTemplate, InlineImage

    # helper that also fetches raw value (extra query)
    def _fetch_raw_value(report: str, var: str) -> str | None:
        q = "SELECT value, anchor_name FROM report_variables WHERE report_name = ? AND var_name = ? ORDER BY created_at DESC LIMIT 1"
        with sqlite3.connect(DB_PATH) as con:
            row = con.execute(q, (report, var)).fetchone()
        return row[0] if row else None

    st.title("📤 Export Final Report")
    reports_df = get_all_reports(DB_PATH)
    if reports_df.empty:
        st.info("No reports available.")
        st.stop()

    chosen_report = st.selectbox("Select Report", reports_df["report_name"].tolist())

    # ---------------- Variable snapshot ----------------
    snap_df = get_variable_status(chosen_report, DB_PATH)
    # if "anchor_name" in snap_df.columns:
    #     snap_df = snap_df[["anchor_name"] + [c for c in snap_df.columns if c != "anchor_name"]]

    def _age_style(r):
        return ["background-color:#ffd6d6" if r.get("age_days", 0) > 5 else "" for _ in r]

    st.markdown("### 🧠 Variables Snapshot")
    st.dataframe(snap_df.style.hide(axis="index").apply(_age_style, axis=1), use_container_width=True)

    # ---------------- FILTERS SECTION ----------------
    st.markdown("### 🔍 Filters")
    
    # Debug: Show available columns
    st.caption(f"Available columns: {', '.join(snap_df.columns.tolist())}")
    
    # Create filter columns
    filter_col1, filter_col2 = st.columns(2)
    
    # Initialize filter variables
    selected_modules = []
    selected_datetime_combinations = []
    module_column = None
    created_column = None
    
    with filter_col1:
        # Module Name Filter - Check multiple possible column names
        possible_module_columns = ['module_name', 'Module', 'module', 'Module_Name']
        
        for col in possible_module_columns:
            if col in snap_df.columns:
                module_column = col
                break
        
        if module_column and not snap_df[module_column].dropna().empty:
            # CASCADING FILTER: Filter by selected dates first (if any dates are selected)
            if 'selected_dates' in locals() and selected_dates and created_column:
                # Only show modules that have variables on the selected dates
                date_filtered_for_modules = snap_df[snap_df['date_only'].isin(selected_dates)]
                module_counts = date_filtered_for_modules[module_column].value_counts().sort_index()
                filter_note = f" (filtered by {len(selected_dates)} selected date(s))"
            else:
                # Show all modules if no date filter is active
                module_counts = snap_df[module_column].value_counts().sort_index()
                filter_note = " (all dates)"
            
            if not module_counts.empty:
                # Create formatted options for display
                module_options = []
                module_mapping = {}
                
                for module_name, count in module_counts.items():
                    if pd.notna(module_name):
                        display_text = f"{module_name} ({count} variables)"
                        module_options.append(display_text)
                        module_mapping[display_text] = module_name
                
                selected_module_displays = st.multiselect(
                    f"📁 Filter by Module(s) - {module_column}{filter_note}",
                    options=module_options,
                    default=module_options,  # Default to all available modules selected
                    help="Select one or more modules to view. Options are filtered by selected date(s)."
                )
                
                # Convert back to actual module names
                selected_modules = [module_mapping[display] for display in selected_module_displays]
            else:
                selected_modules = []
                st.info("No modules found for the selected date(s)")
        else:
            available_cols = [col for col in snap_df.columns if 'module' in col.lower()]
            if available_cols:
                st.info(f"Module columns found but empty: {', '.join(available_cols)}")
            else:
                st.info("No module name column found in data")
    
    with filter_col2:
        # Split Date/Time Filter Approach
        possible_date_columns = ['created_at', 'Created_At', 'created', 'timestamp', 'date']
        
        for col in possible_date_columns:
            if col in snap_df.columns:
                created_column = col
                break
        
        if created_column and not snap_df[created_column].dropna().empty:
            # Convert to datetime
            snap_df[created_column] = pd.to_datetime(snap_df[created_column])
            
            # Create separate day and time columns
            snap_df['date_only'] = snap_df[created_column].dt.date
            snap_df['time_only'] = snap_df[created_column].dt.time
            
            # Create sub-columns for day and time
            date_subcol1, date_subcol2 = st.columns(2)
            
            with date_subcol1:
                # Day Filter
                date_counts = snap_df['date_only'].value_counts().sort_index(ascending=False)
                
                if not date_counts.empty:
                    date_options = []
                    date_mapping = {}
                    
                    for date_val, count in date_counts.items():
                        if pd.notna(date_val):
                            display_text = f"{date_val} ({count})"
                            date_options.append(display_text)
                            date_mapping[display_text] = date_val
                    
                    # Default to most recent date
                    default_date = [date_options[0]] if date_options else []
                    
                    selected_date_displays = st.multiselect(
                        "📅 Day",
                        options=date_options,
                        default=default_date,
                        help="Select day(s)"
                    )
                    
                    selected_dates = [date_mapping[display] for display in selected_date_displays if display in date_mapping]
                else:
                    selected_dates = []
            
            with date_subcol2:
                # Time Filter (only show times for selected dates)
                if selected_dates:
                    # Filter data by selected dates first
                    date_filtered_df = snap_df[snap_df['date_only'].isin(selected_dates)]
                    time_counts = date_filtered_df['time_only'].value_counts().sort_index(ascending=False)
                    
                    if not time_counts.empty:
                        time_options = []
                        time_mapping = {}
                        
                        for time_val, count in time_counts.items():
                            if pd.notna(time_val):
                                display_text = f"{time_val.strftime('%H:%M:%S')} ({count})"
                                time_options.append(display_text)
                                time_mapping[display_text] = time_val
                        
                        # Default to all times for selected dates
                        selected_time_displays = st.multiselect(
                            "🕐 Time",
                            options=time_options,
                            default=time_options,  # All times by default
                            help="Select time(s) for chosen day(s)"
                        )
                        
                        selected_times = [time_mapping[display] for display in selected_time_displays if display in time_mapping]
                    else:
                        selected_times = []
                        st.info("No times found for selected dates")
                else:
                    selected_times = []
                    st.info("Select a date first")
            
            # Combine date and time filters
            if selected_dates and selected_times:
                # Create datetime combinations for filtering
                selected_datetime_combinations = []
                for date_val in selected_dates:
                    for time_val in selected_times:
                        # Find rows that match both date and time
                        matching_rows = snap_df[
                            (snap_df['date_only'] == date_val) & 
                            (snap_df['time_only'] == time_val)
                        ]
                        if not matching_rows.empty:
                            selected_datetime_combinations.extend(matching_rows[created_column].tolist())
        else:
            available_cols = [col for col in snap_df.columns if any(term in col.lower() for term in ['date', 'time', 'created'])]
            if available_cols:
                st.info(f"Date columns found but empty: {', '.join(available_cols)}")
            else:
                st.info("No date column found in data")

    # Apply filters to snap_df
    filtered_snap_df = snap_df.copy()
    
    # Apply module filter
    if selected_modules and module_column:
        filtered_snap_df = filtered_snap_df[filtered_snap_df[module_column].isin(selected_modules)]
    
    # Apply date/time filter
    if selected_datetime_combinations and created_column:
        filtered_snap_df = filtered_snap_df[filtered_snap_df[created_column].isin(selected_datetime_combinations)]
    
    # Show filter results
    total_modules = len(snap_df[module_column].dropna().unique()) if module_column else 0
    total_datetimes = len(snap_df[created_column].dropna().unique()) if created_column else 0
    
    # Check if filters are active
    module_filter_active = len(selected_modules) < total_modules if total_modules > 0 else False
    datetime_filter_active = len(selected_datetime_combinations) < total_datetimes if total_datetimes > 0 else False
    
    if len(filtered_snap_df) != len(snap_df) or module_filter_active or datetime_filter_active:
        filter_info_col1, filter_info_col2 = st.columns(2)
        with filter_info_col1:
            st.info(f"🔍 **Filter Results:** {len(filtered_snap_df)} of {len(snap_df)} variables")
            
            # Show active filters
            active_filters = []
            if module_filter_active:
                active_filters.append(f"📁 {len(selected_modules)} of {total_modules} modules")
            if datetime_filter_active:
                active_filters.append(f"📅 {len(selected_datetime_combinations)} of {total_datetimes} records")
            
            if active_filters:
                st.caption("Active filters: " + " • ".join(active_filters))
        
        with filter_info_col2:
            if st.button("🔄 Reset All Filters", help="Reset to show all variables"):
                st.rerun()
    
    # Display filtered summary if filters are applied
    if len(filtered_snap_df) < len(snap_df):
        with st.expander("📊 Filtered Variables Summary", expanded=False):
            st.dataframe(
                filtered_snap_df.style.hide(axis="index").apply(_age_style, axis=1), 
                use_container_width=True
            )

    # ---------------- Visualisation chooser -------------
    st.markdown("### 📊 Select Tables/Charts to Visualize")
    
    # Use filtered data for available options
    available = filtered_snap_df.get("var_name", pd.Series(dtype=str)).tolist()
    
    if not available:
        st.warning("⚠️ No variables available with current filters. Please adjust your filter criteria.")
    else:
        # Group options by module for better UX
        if module_column and module_column in filtered_snap_df.columns:
            # Create a dictionary of module -> variables
            module_vars = {}
            for _, row in filtered_snap_df.iterrows():
                module = row.get(module_column, "Unknown")
                var_name = row.get("var_name", "")
                if module not in module_vars:
                    module_vars[module] = []
                if var_name:
                    module_vars[module].append(var_name)
            
            # Show module breakdown
            with st.expander(f"📁 Available Variables by Module ({len(available)} total)", expanded=False):
                for module, vars_list in module_vars.items():
                    st.write(f"**{module}:** {len(vars_list)} variables")
                    st.write(", ".join(vars_list[:5]) + ("..." if len(vars_list) > 5 else ""))
        
        selected_tables = st.multiselect(
            "Choose table(s)/chart(s) to visualize", 
            available, 
            placeholder="Start typing variable names..."
        )

        if selected_tables:
            st.markdown("### 📊 GreatTables Images & Raw Values")
            for var_name in selected_tables:
                # Get variable info for better display
                var_rows = filtered_snap_df[filtered_snap_df["var_name"] == var_name]
                if not var_rows.empty:
                    var_info = var_rows.iloc[0]
                    module_name = var_info.get(module_column, "Unknown") if module_column else "Unknown"
                    created_at = var_info.get(created_column, "Unknown") if created_column else "Unknown"
                else:
                    module_name = "Unknown"
                    created_at = "Unknown"
                
                # Display variable metadata
                st.markdown(f"#### 📈 {var_name}")
                meta_col1, meta_col2 = st.columns(2)
                with meta_col1:
                    st.markdown(f"**Module:** `{module_name}`")
                with meta_col2:
                    st.markdown(f"**Created:** `{created_at}`")
                
                # Fetch gt_image (bytes or path) and anchor_name
                gt_image, anchor_name = fetch_gt_image(chosen_report, var_name, DB_PATH)
                if gt_image:
                    st.write(f"🔍 Image length: {len(str(gt_image))} bytes")
                    try:
                        if isinstance(gt_image, str):  # If gt_image is a file path (for charts)
                            image_path = Path(gt_image)
                            if image_path.exists():
                                image = Image.open(image_path)
                                st.image(image, caption=f"Image for {anchor_name}", use_container_width=True)
                            else:
                                st.warning(f"Image file not found at {gt_image} for {var_name}")
                        else:  # If gt_image is bytes (for great_tables)
                            image = Image.open(io.BytesIO(gt_image))
                            st.image(image, caption=f"Image for {anchor_name}", use_container_width=True)
                    except Exception as exc:
                        st.warning(f"Failed to display image for {var_name}: {exc}")
                else:
                    st.warning(f"No image found for {var_name}")

                st.markdown(f"**Anchor name:** `{anchor_name or '–'}`")
                st.markdown(f"**Raw value for `{var_name}`:**")
                _pretty_print_value(_fetch_raw_value(chosen_report, var_name))
                
                st.markdown("---")  # Separator between variables

    # ---------------- Template render block -------------
    st.markdown("### 📄 Select Template or Use Existing Partial")
    tmpl_dir = Path("reporting/templates/docx")
    tmpl_files = list(tmpl_dir.glob("*.docx"))
    if not tmpl_files:
        st.error("No .docx templates found in reporting/templates/docx")
        st.stop()

    tmpl_choice = st.selectbox("Choose a template:", [p.name for p in tmpl_files])
    tmpl_path = tmpl_dir / tmpl_choice

    def _build_context(report: str, tpl: DocxTemplate) -> dict:
        q = "SELECT anchor_name, value, gt_image FROM report_variables WHERE report_name = ? ORDER BY created_at"
        df = pd.read_sql_query(q, sqlite3.connect(DB_PATH), params=(report,))
        ctx: dict[str, object] = {}
        for _idx, row in df.iterrows():
            anchor = row["anchor_name"]
            if row["gt_image"]:
                if isinstance(row["gt_image"], str):  # If gt_image is a file path (for charts)
                    image_path = Path(row["gt_image"])
                    if image_path.exists():
                        with open(image_path, "rb") as f:
                            image_bytes = f.read()
                        ctx[anchor] = InlineImage(tpl, io.BytesIO(image_bytes), width=docx.shared.Inches(5))
                    else:
                        st.warning(f"Image file not found at {row['gt_image']} for {anchor}")
                        ctx[anchor] = None  # Skip in template to avoid breaking
                else:  # If gt_image is bytes (for great_tables)
                    ctx[anchor] = InlineImage(tpl, io.BytesIO(row["gt_image"]), width=docx.shared.Inches(5))
            else:
                try:
                    ctx[anchor] = json.loads(row["value"])
                except (TypeError, ValueError):
                    ctx[anchor] = row["value"]
        return ctx
    
    if st.button("📄 Render Final Report"):
        tpl = DocxTemplate(str(tmpl_path))
        context = _build_context(chosen_report, tpl)

        missing = tpl.get_undeclared_template_variables() - set(context.keys())
        if missing:
            st.warning("⚠️ Missing anchors: " + ", ".join(missing))
        else:
            st.success("All template anchors matched!")

        tpl.render(context, jinja_env=Environment(undefined=DebugUndefined))
        out_dir = Path("app_files") / chosen_report / datetime.today().strftime("%Y-%m-%d")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / "Final_Report.docx"
        tpl.save(out_path)
        st.success(f"Report saved → {out_path.absolute()}")
        with open(out_path, "rb") as fh:
            st.download_button("📥 Download Final Report", fh.read(), file_name="Final_Report.docx")

    # staged_docx save (unchanged)
    if "staged_docx" in st.session_state:
        with st.expander("💾 Save or Preview Staged Report", expanded=True):
            fname = f"partial_{chosen_report}_{datetime.today().strftime('%Y-%m-%d')}.docx"
            if st.button("💾 Save to app_files"):
                out = Path("app_files") / fname
                st.session_state.staged_docx.save(out)
                st.success(f"Saved partial report as {fname} in app_files/")
                
#--------------------------------------------------------------------------
# --- SINGLE UPLOAD -----------------------------------
#--------------------------------------------------------------------------
elif selected_section == "single_upload":
    st.title("📂 Single File Upload")
    # Load reports from DB
    reports_df = get_all_reports(DB_PATH)
    report_names = ["-- Create new --"] + reports_df["report_name"].tolist()

    # Default selection logic
    default_report_name = (
        st.session_state.selected_report_after_create
        if st.session_state.selected_report_after_create in report_names
        else report_names[0]
    )

    chosen_report = st.selectbox("Choose a report:", report_names, index=report_names.index(default_report_name))
    if chosen_report == "-- Create new --":
        st.markdown("### Create New Report")
        new_report_name = st.text_input("Enter the name for the new report:", key="new_report_name_input")

        if st.button("➕ Create Report", key="create_report_button"):
            if new_report_name.strip():
                try:
                    # Create the report in the database
                    create_new_report(new_report_name.strip(), DB_PATH)

                    # Store the new report name in session state for auto-selection after rerun
                    st.session_state.selected_report_after_create = new_report_name.strip()

                    # Show feedback and rerun
                    st.success(f"Report '{new_report_name.strip()}' created!")
                    st.toast(f"Report '{new_report_name.strip()}' created!", icon="✅") # Use toast for notification across rerun
                    st.rerun()

                except ValueError as e:
                    st.error(str(e)) # Report specific database error (e.g., name already exists)
                except Exception as e:
                    st.error(f"An unexpected error occurred during report creation: {e}")
                    print(traceback.format_exc()) # Print traceback

            else:
                st.warning("Please provide a name for the new report.")

        # Use st.stop() to prevent the file upload UI from showing when in "Create new" mode
        st.stop()

    else: # A report *other* than "-- Create new --" is selected
            st.markdown(f"### Upload file for report: `{chosen_report}`")

            # File uploader
            uploaded_file = st.file_uploader("📁 Upload .xlsx or .csv file", type=["xlsx", "xls", "csv"], key=f"single_upload_file_uploader_{st.session_state.file_uploader_key_counter}")

            # Stop here if no file is uploaded yet, otherwise proceed with file processing
            if not uploaded_file:
                st.info("Awaiting file upload...")
                st.stop() # Keep st.stop() here

            # --- File Processing and Transformation Logic ---
            filename = uploaded_file.name
            extension = os.path.splitext(filename)[1].lower()
            filename_wo_ext = os.path.splitext(filename)[0]
            # Derive a raw table name (used for the 'table_name' column in upload_log)
            default_raw_table_name = f"raw_{filename_wo_ext.lower()}"

            st.success(f"📥 File received: `{filename}`")
            st.info(f"Linking upload to report: `{chosen_report}`")


            # --- Sheet Selection ---
            existing_sheet, saved_start_row = get_existing_rule(filename, DB_PATH)
            sheet_to_use = None
            start_row = 0

            if extension in [".xlsx", ".xls"]:
                try:
                    uploaded_file.seek(0)
                    xls = pd.ExcelFile(uploaded_file)
                    sheet_names = xls.sheet_names

                    if existing_sheet and existing_sheet in sheet_names:
                        st.success(f"✅ Found saved sheet: `{existing_sheet}`")
                        use_existing = st.checkbox("Use saved sheet rule?", value=True, key="use_existing_sheet_checkbox")
                        sheet_to_use = existing_sheet if use_existing else None
                        start_row = saved_start_row if (use_existing and saved_start_row is not None) else 0

                    if sheet_to_use is None:
                        sheet_to_use = st.selectbox("📑 Select sheet to upload:", sheet_names, key="sheet_select_box")
                        start_row = st.number_input("Start row (0-indexed header):", min_value=0, value=0, key="sheet_start_row_input")
                    else:
                        start_row = st.number_input("Start row (0-indexed header):", min_value=0, value=saved_start_row or 0, key="sheet_start_row_input_existing")

                    if not sheet_to_use:
                        st.warning("Please select a sheet.")
                        st.stop()

                    if st.button("💾 Save sheet rule", key="save_sheet_rule_button"):
                        insert_sheet_rule(filename, sheet_to_use, start_row, DB_PATH)
                        st.success(f"Rule saved for `{filename}`: sheet `{sheet_to_use}`, header row {start_row + 1} (0-indexed row {start_row})")
                        st.rerun()

                except Exception as excel_error:
                    st.error(f"Error reading Excel file or selecting sheet: {excel_error}")
                    print(traceback.format_exc())
                    st.stop()

            elif extension == ".csv":
                sheet_to_use = "CSV_SHEET"
                start_row = st.number_input("Start row (0-indexed header):", min_value=0, value=saved_start_row or 0, key="csv_start_row_input")

                if st.button("💾 Save sheet rule", key="save_csv_sheet_rule_button"):
                    insert_sheet_rule(filename, sheet_to_use, start_row, DB_PATH)
                    st.success(f"Rule saved for `{filename}`: CSV, header row {start_row + 1} (0-indexed row {start_row})")
                    st.rerun()

            else:
                st.error("❌ Unsupported file format.")
                st.stop()


            # --- Load Preview DataFrame based on Sheet/Row Rules (Common Block) ---
            preview_df = None
            try:
                if extension in [".xlsx", ".xls"] and sheet_to_use:
                    uploaded_file.seek(0)
                    xls = pd.ExcelFile(uploaded_file)
                    preview_df = xls.parse(sheet_to_use, skiprows=start_row)
                elif extension == ".csv":
                    uploaded_file.seek(0)
                    preview_df = pd.read_csv(uploaded_file, skiprows=start_row)
            except Exception as load_preview_error:
                st.error(f"Error loading data preview from file: {load_preview_error}")
                print(traceback.format_exc())
                st.stop()


            if preview_df is None or preview_df.empty:
                st.warning("Could not load data preview or data is empty after applying start row.")
                st.stop()


            # --- Transformation Rules ---
            rules = get_transform_rules(filename, sheet_to_use, DB_PATH)
            current_columns = preview_df.columns.tolist()

            if not rules:
                st.info("No saved transformation rules found for this file/sheet. Proposing initial rules based on preview columns.")
                rules = [
                    {"original_column": col, "renamed_column": col, "included": True, "filename": filename, "sheet": sheet_to_use, "created_at": datetime.now().isoformat()}
                    for col in current_columns
                ]

            st.markdown("### Column Transformations")
            st.write("Review and rename/exclude columns. Changes are saved upon Upload.")

            cols_header = st.columns([0.4, 0.4, 0.2])
            with cols_header[0]: st.write("**Renamed Column**")
            with cols_header[1]: st.write("**Original Column**")
            with cols_header[2]: st.write("**Include?**")
            st.markdown("---")

            edited_rules = []
            for rule in rules:
                original_col = rule.get("original_column", "Unknown Column")

                if original_col in preview_df.columns:
                    widget_key_prefix = f"{filename}_{sheet_to_use}_{original_col}"

                    col_r, col_o, col_i = st.columns([0.4, 0.4, 0.2])
                    with col_r:
                        renamed_col = st.text_input(f"Rename {original_col}", value=rule.get("renamed_column", original_col), label_visibility="collapsed", key=f"rename_{widget_key_prefix}")
                    with col_o:
                        st.text_input(f"Original {original_col}", value=original_col, disabled=True, label_visibility="collapsed", key=f"original_{widget_key_prefix}")
                    with col_i:
                        included = st.checkbox(f"Include {original_col}", value=rule.get("included", True), label_visibility="collapsed", key=f"include_{widget_key_prefix}")

                    edited_rules.append({
                        "filename": filename,
                        "sheet": sheet_to_use,
                        "original_column": original_col,
                        "renamed_column": renamed_col,
                        "included": included,
                        "created_at": rule.get("created_at", datetime.now().isoformat())
                    })


            # --- Determine Default / Proposed Database Table Name ---
            # Derive a default name from the filename (cleaned)
            default_derived_table_name = filename_wo_ext.lower().replace(" ", "_").replace("-", "_").replace(".", "_")
            import re
            default_derived_table_name = re.sub(r'\W+', '_', default_derived_table_name) # Replace non-alphanumeric with _
            default_derived_table_name = re.sub(r'^_+', '', default_derived_table_name) # Remove leading underscores from cleaning
            if not default_derived_table_name:
                default_derived_table_name = "uploaded_data" # Fallback if cleaning results in empty string


            # --- Database Table Naming Section ---
            st.markdown("### Database Table Naming")
            st.write("Specify the name for the database table where this file's data will be stored.")
            st.write(f"Proposed default name based on filename: `{default_derived_table_name}`")


            # Provide a text input defaulted to the derived name
            user_defined_table_name = st.text_input(
                "Database Table Name:",
                value=default_derived_table_name, # Default to the derived name
                help="Must be unique and contain only letters, numbers, and underscores.",
                key="database_table_name_input"
            )

            # Validate the user's input name
            cleaned_input_name = user_defined_table_name.strip()
            # More robust validation pattern
            is_valid_name = bool(re.fullmatch(r'[a-zA-Z_][a-zA-Z0-9_]*', cleaned_input_name))

            if not cleaned_input_name:
                st.warning("Table name cannot be empty.")
                is_valid_name = False
            elif not is_valid_name:
                st.warning("Table name must start with a letter or underscore and contain only letters, numbers, or underscores.")


            # --- Transformed Data Preview ---
            # Apply rules to preview and display it
            cols_to_include_original = [r["original_column"] for r in edited_rules if r["included"] and r["original_column"] in preview_df.columns]
            preview_rename_map = {r["original_column"]: r["renamed_column"] for r in edited_rules if r["included"] and r["original_column"] in preview_df.columns}

            final_preview_df = pd.DataFrame() # Initialize
            if cols_to_include_original: # Only proceed if there are columns to include
                try:
                    final_preview_df = preview_df[cols_to_include_original].copy()
                    final_preview_df.rename(columns=preview_rename_map, inplace=True)
                    st.info("✅ Preview updated based on transformations.")
                except KeyError as e:
                    st.error(f"Error applying transformations for preview: Column missing or misspelled after rules applied: {e}. Showing empty preview.")
                    print(traceback.format_exc())
                except Exception as e:
                    st.error(f"An unexpected error occurred applying transformations for preview: {e}. Showing empty preview.")
                    print(traceback.format_exc())
            else:
                st.warning("No columns selected for inclusion after transformations.")


            st.markdown("### 👀 Transformed Data Preview (First 10 rows)")
            if final_preview_df.empty:
                st.info("Preview is empty or no columns selected for display.")
            else:
                st.dataframe(final_preview_df.head(10))


            # --- Save Rules & Upload Button ---
            st.markdown("---")
            # The button should only be enabled if a valid table name is provided
            if st.button("✅ Save Rules & Upload Data", key="save_rules_and_upload_button", disabled=not is_valid_name):
                # This code block only runs when the button is clicked AND the name is valid

                st.info("Starting data upload process...")
                now = datetime.now().isoformat()

                # --- Save Edited Transform Rules ---
                for rule in edited_rules:
                    rule["created_at"] = now
                save_transform_rules(edited_rules, DB_PATH)
                st.success("✅ Transformation rules saved.")


                # --- Apply Rules to FULL DataFrame and Upload ---
                full_df = None
                try:
                    uploaded_file.seek(0)
                    if extension in [".xlsx", ".xls"]:
                        xls = pd.ExcelFile(uploaded_file)
                        full_df = xls.parse(sheet_to_use, skiprows=start_row)
                    elif extension == ".csv":
                        full_df = pd.read_csv(uploaded_file, skiprows=start_row)
                    st.write("🧪 DEBUG: Full DataFrame shape:", full_df.shape)
                except Exception as load_full_error:
                    st.error(f"Error loading full data from file: {load_full_error}")
                    print(traceback.format_exc())
                    # Don't stop here, let the error message show, the check below handles full_df being None
                    full_df = None # Ensure full_df is None if loading fails

                # --- Validate + Transform ---
                if full_df is not None and not full_df.empty:
                    try:
                        # Filter and rename the *full* DataFrame based on the saved/edited rules
                        cols_to_include_original_full = [r["original_column"] for r in edited_rules if r["included"] and r["original_column"] in full_df.columns]
                        upload_df = full_df[cols_to_include_original_full].copy()
                        upload_rename_map = {r["original_column"]: r["renamed_column"] for r in edited_rules if r["included"] and r["original_column"] in full_df.columns}
                        upload_df.rename(columns=upload_rename_map, inplace=True)

                        st.write("🧪 DEBUG: Upload DF shape:", upload_df.shape)
                        st.write("🧪 DEBUG: Upload DF columns:", upload_df.columns.tolist())

                        if upload_df.empty:
                            st.warning("Upload resulted in an empty dataset after applying transformations. No data uploaded.")
                            # Don't stop here, let the warning show, the rest of the block handles upload_df being empty
                            # exit the try block gracefully
                            raise ValueError("Upload resulted in empty data after transformations.") # Raise to jump to except


                        # --- Final Alias & Database Save using User-Defined Name ---

                        # Use the user-defined name as the actual table name
                        final_table_name_in_db = cleaned_input_name # Use the validated name from the text input
                        st.write("🧪 DEBUG: Final table name:", final_table_name_in_db)

                        # Register the user's chosen table name AS the alias for this filename.
                        # This makes the alias match the table name for simplicity.
                        file_alias = final_table_name_in_db
                        register_file_alias(filename, file_alias, db_path=DB_PATH)
                        st.info(f"Alias `{file_alias}` registered for file `{filename}`.")

                        with sqlite3.connect(DB_PATH) as conn:
                            # Insert upload log using the original filename, the derived *raw* table_name (for log detail),
                            # the selected report, AND the user-defined table name AS the alias
                            upload_id = insert_upload_log(
                                filename, 
                                default_raw_table_name, 
                                upload_df.shape[0], 
                                upload_df.shape[1],
                                chosen_report, 
                                table_alias=file_alias, 
                                db_path=DB_PATH # Log the alias (user's name)
                            )
                            st.success(f"✅ Upload log created (ID: {upload_id}).")
                            # Append metadata
                            upload_df["upload_id"] = upload_id
                            upload_df["uploaded_at"] = now
                            
                            st.write("🧪 DEBUG: Saving table", final_table_name_in_db)
                            st.write("🧪 DEBUG: Upload DF preview", upload_df.head(3))

                            # Save to SQL
                            upload_df.to_sql(final_table_name_in_db, conn, index=False, if_exists="replace")
                            conn.commit()
                            sample_df = pd.read_sql_query(f"SELECT * FROM `{final_table_name_in_db}` LIMIT 5", conn)
                            st.markdown(f"### 🧪 Sample of `{final_table_name_in_db}` from DB")
                            st.dataframe(sample_df)

                            # Only show success if data is there
                            if not sample_df.empty:
                                st.success(f"📦 Uploaded to table `{final_table_name_in_db}` with {len(sample_df)}+ rows visible.")
                            else:
                                st.warning(f"⚠️ Upload to `{final_table_name_in_db}` completed, but no data appears in preview. Check start row / rules.")

                            # Verify insert
                            try:
                                result = pd.read_sql_query(f"SELECT COUNT(*) AS cnt FROM `{final_table_name_in_db}`", conn)
                                st.write(f"🧪 DEBUG: Row count in `{final_table_name_in_db}` after insert:", result['cnt'].iloc[0])
                            except Exception as verify_error:
                                st.error(f"⚠️ Failed to verify row count: {verify_error}")

                    except ValueError as ve: # Catch the ValueError for empty data gracefully
                        st.warning(f"Upload failed: {ve}")
                    except Exception as e:
                        st.error(f"❌ Upload failed: {e}")
                        import traceback
                        st.code(traceback.format_exc())

                    # 🔄 Trigger reset
                    st.toast(f"✅ Upload complete for `{filename}` → `{final_table_name_in_db}`", icon="📥")
                    st.session_state.file_uploader_key_counter += 1
                    st.rerun()

                elif full_df is not None and full_df.empty:
                    st.warning("Upload failed: Loaded data from file is empty after applying start row.")

                else: # full_df is None
                    st.error("❌ Upload failed: Could not load data from file.")


# ------------------------------------------------------------------
# 4) MASS UPLOAD
# ------------------------------------------------------------------
elif selected_section == "mass_upload":
    st.subheader("📦 Mass Upload")

    # ─────────────────────────── 1. drag-and-drop ──────────────────────────
    uploads = st.file_uploader(
        "Drag & drop .xlsx / .xls / .csv files here, or click to select",
        type=["xlsx", "xls", "csv"],
        accept_multiple_files=True
    )
    for uf in uploads:
        dest = Path("app_files") / uf.name
        dest.write_bytes(uf.getbuffer())
        st.success(f"Saved **{uf.name}** to `app_files/`")

    # ─────────────────────────── 2. choose report ──────────────────────────
    reports_df = get_all_reports(DB_PATH)
    if reports_df.empty:
        st.warning("No reports yet – create one in *Single File Upload*.")
        st.stop()

    chosen_report = st.selectbox("Report:", reports_df["report_name"])

    all_files = [f for f in os.listdir("app_files")
                 if f.lower().endswith((".csv", ".xlsx", ".xls"))]
    if not all_files:
        st.info("No files found in `app_files/`.")
        st.stop()

    # ─────────────────────────── 3. pre-flight checks ──────────────────────
    ready_files, not_ready = [], []

    for file in all_files:
        alias = get_alias_for_file(file, DB_PATH)
        sheet, start_row = get_existing_rule_for_report(chosen_report, file, DB_PATH)

        if alias and alias_exists(alias, DB_PATH) and sheet:
            ready_files.append((file, alias, sheet, start_row or 0))
        else:
            not_ready.append(file)

    # display status
    if ready_files:
        st.success("Files ready for upload: " + ", ".join(f"`{f[0]}`" for f in ready_files))
    if not_ready:
        st.warning(
            "Files **not ready** (missing alias or loading rule): "
            + ", ".join(f"`{f}`" for f in not_ready)
            + ".\nOpen **Single File Upload** to create the missing alias/rule."
        )

    # ─────────────────────────── 4. upload button ──────────────────────────
    if ready_files:
        if st.button("🚀 Upload All Ready Files"):
            ok_cnt, fail_cnt = 0, 0
            with st.spinner("Uploading…"):
                for file, alias, sheet, start_row in ready_files:
                    fp = Path("app_files") / file
                    ext = fp.suffix.lower()

                    # 4a read file
                    try:
                        if ext in {".xlsx", ".xls"}:
                            df = pd.read_excel(fp, sheet_name=sheet, skiprows=start_row)
                        else:
                            df = pd.read_csv(fp, skiprows=start_row)
                    except Exception as e:
                        st.error(f"❌ **{file}** – read error: {e}")
                        fail_cnt += 1
                        continue

                    # 4b transform rules (if any)
                    rules = get_transform_rules(file, sheet, DB_PATH)
                    if rules:
                        try:
                            keep = [r["original_column"] for r in rules
                                    if r["included"] and r["original_column"] in df.columns]
                            rename = {r["original_column"]: r["renamed_column"]
                                      for r in rules if r["included"] and r["original_column"] in df.columns}
                            df = df[keep].rename(columns=rename)
                        except Exception as e:
                            st.error(f"❌ **{file}** – transform error: {e}")
                            fail_cnt += 1
                            continue

                    # 4c write to DB
                    try:
                        upload_id = insert_upload_log(
                            file, f"raw_{fp.stem.lower()}",
                            df.shape[0], df.shape[1], chosen_report,
                            table_alias=alias, db_path=DB_PATH
                        )
                        df["upload_id"] = upload_id
                        df["uploaded_at"] = datetime.now().isoformat()

                        with sqlite3.connect(DB_PATH) as con:
                            df.to_sql(alias, con, if_exists="replace", index=False)

                        update_alias_status(alias, file, DB_PATH)
                        st.success(f"✅ {file} → table **{alias}**")
                        ok_cnt += 1
                    except Exception as e:
                        st.error(f"❌ **{file}** – DB error: {e}")
                        traceback.print_exc()
                        fail_cnt += 1

            st.info(f"Upload summary – {ok_cnt} OK · {fail_cnt} failed")
    else:
        st.info("No file is fully configured yet, so the upload button is hidden.")

# ------------------------------------------------------------------
# 5) HISTORY
# ------------------------------------------------------------------
elif selected_section == "history":
    # --- Section: Upload History & Management ---
    print("DEBUG: Entering history section")
    try:
        st.subheader("🔎 Upload History & Management")

        # --- Confirmation Step for Upload Deletion (appears when state is set) ---
        if st.session_state.pending_delete_uploads is not None:
            upload_ids_to_confirm_delete = st.session_state.pending_delete_uploads
            num_uploads_to_delete = len(upload_ids_to_confirm_delete)

            st.error("🛑 **Critical Warning: Confirm Upload Deletion** 🛑")
            st.markdown(f"""
            You are about to delete the following **{num_uploads_to_delete} specific upload record(s)**:
            """)

            # Fetch details for the uploads being deleted to display in the warning
            try:
                with sqlite3.connect(DB_PATH) as conn:
                    placeholders = ','.join('?' for _ in upload_ids_to_confirm_delete)
                    query = f"""
                        SELECT filename, table_alias, uploaded_at, rows, cols, id
                        FROM upload_log
                        WHERE id IN ({placeholders})
                        ORDER BY uploaded_at DESC
                    """
                    uploads_to_delete_df = pd.read_sql_query(query, conn, params=upload_ids_to_confirm_delete)
                    uploads_to_delete_df.rename(columns={'id': 'upload_id'}, inplace=True)

                if not uploads_to_delete_df.empty:
                    st.dataframe(uploads_to_delete_df)
                else:
                    st.warning("Details for selected uploads could not be retrieved.")

            except Exception as fetch_error:
                 st.error(f"Error fetching upload details for confirmation: {fetch_error}")
                 print(traceback.format_exc())


            st.markdown("""
            **Deleting an upload record will:**
            * **Permanently remove the history entry** for this specific file upload from the logs.
            * **Permanently delete the data rows from the corresponding data table** (e.g., the table named after the alias) that belong to this specific upload (`WHERE upload_id = this_id`).
            * **This action CANNOT be undone.**
            """)

            col_confirm_upload, col_cancel_upload = st.columns(2)

            with col_confirm_upload:
                if st.button(f"✅ **Confirm Delete {num_uploads_to_delete} Upload Record(s) AND Data**", key="confirm_execute_delete_uploads"):
                    st.info("Executing upload and data deletion...")
                    deleted_count = 0
                    uploads_list_to_delete = st.session_state.pending_delete_uploads

                    with sqlite3.connect(DB_PATH) as conn:
                         cursor = conn.cursor()
                         for upload_id in uploads_list_to_delete:
                             try:
                                 cursor.execute("SELECT table_alias FROM upload_log WHERE id = ?", (upload_id,))
                                 row = cursor.fetchone()
                                 if row and row[0]:
                                     table_to_clean = row[0]

                                     delete_data_sql = f"DELETE FROM {table_to_clean} WHERE upload_id = ?"
                                     try:
                                        cursor.execute(delete_data_sql, (upload_id,))
                                        st.write(f"✅ Deleted data from `{table_to_clean}` for upload ID `{upload_id}`")
                                     except sqlite3.OperationalError as oe:
                                         st.warning(f"⚠️ Data table `{table_to_clean}` or upload_id column not found for upload ID `{upload_id}`. Skipping data deletion for this upload.")
                                     except Exception as data_delete_e:
                                         st.error(f"❌ Error deleting data for upload ID `{upload_id}` from `{table_to_clean}`: {data_delete_e}")
                                         print(traceback.format_exc())

                                 cursor.execute("DELETE FROM upload_log WHERE id = ?", (upload_id,))
                                 st.write(f"✅ Deleted log entry for upload ID `{upload_id}`")
                                 deleted_count += 1

                             except Exception as e:
                                 st.error(f"❌ Failed to delete upload record for ID `{upload_id}`: {e}")
                                 print(traceback.format_exc())

                         conn.commit()
                    st.session_state.pending_delete_uploads = None
                    st.success(f"Completed deleting {deleted_count} upload record(s) and associated data.")
                    st.toast(f"Deleted {deleted_count} uploads!", icon="🗑️")
                    st.rerun()

            with col_cancel_upload:
                if st.button("↩️ Cancel Upload Deletion", key="cancel_delete_uploads"):
                    st.session_state.pending_delete_uploads = None
                    st.info("Upload deletion cancelled.")
                    st.rerun()

            st.stop()


        # --- Confirmation Step for Metadata Deletion (appears when state is set, secondary) ---
        elif st.session_state.pending_delete_reports is not None:
            reports_to_confirm_delete = st.session_state.pending_delete_reports
            num_reports_to_delete = len(reports_to_confirm_delete)

            st.error("🛑 **Critical Warning: Confirm Metadata Deletion** 🛑")
            st.markdown(f"""
            You are about to delete the metadata for the following **{num_reports_to_delete} report(s)**:
            """)
            for report_name in reports_to_confirm_delete:
                st.write(f"- `{report_name}`")

            st.markdown("""
            **Deleting metadata will:**
            * Remove the report definition itself from the app's list.
            * **Permanently delete all upload history records associated with these specific reports.** (You will lose information about which files were uploaded when *for this report*).
            * Delete any defined expected structure or cutoff logs linked to these reports.
            * **Effectively "reset" the application's knowledge and validation logic specifically for these reports.**

            **This action does NOT delete the actual data tables** (like tables named after aliases) in the database that contain the content of the uploaded files.
            """)

            col_confirm_meta, col_cancel_meta = st.columns(2)

            with col_confirm_meta:
                if st.button(f"✅ **Confirm Delete Metadata for {num_reports_to_delete} Report(s)**", key="confirm_execute_multi_delete"):
                    st.info("Executing deletion...")
                    deleted_count = 0
                    reports_list_to_delete = st.session_state.pending_delete_reports

                    with sqlite3.connect(DB_PATH) as conn:
                         cursor = conn.cursor()
                         for report_name in reports_list_to_delete:
                             try:
                                 cursor.execute("SELECT COUNT(*) FROM reports WHERE report_name = ?", (report_name,))
                                 if cursor.fetchone()[0] > 0:
                                     cursor.execute("DELETE FROM reports WHERE report_name = ?", (report_name,))
                                     cursor.execute("DELETE FROM upload_log WHERE report_name = ?", (report_name,))
                                     cursor.execute("DELETE FROM report_structure WHERE report_name = ?", (report_name,))
                                     cursor.execute("DELETE FROM report_cutoff_log WHERE report_name = ?", (report_name,))
                                     deleted_count += 1
                                     st.write(f"✅ Deleted metadata for '{report_name}'")
                                 else:
                                     st.write(f"ℹ️ Report '{report_name}' not found, skipping deletion.")

                             except Exception as e:
                                 st.error(f"❌ Failed to delete metadata for '{report_name}': {e}")
                                 print(traceback.format_exc())

                         conn.commit()
                    st.session_state.pending_delete_reports = None
                    st.success(f"Completed deleting metadata for {deleted_count} report(s).")
                    st.toast(f"Metadata deleted for {deleted_count} report(s)!", icon="🗑️")
                    st.rerun()

            with col_cancel_meta:
                if st.button("↩️ Cancel Deletion", key="cancel_multi_delete"):
                    st.session_state.pending_delete_reports = None
                    st.info("Metadata deletion cancelled.")
                    st.rerun()

            st.stop()


        # --- Normal History View (if no confirmation pending) ---
        else:
            # --- Multi-Select Report Metadata Deletion (checkboxes) ---
            st.markdown("### Delete Report Metadata")
            reports_df_all = get_all_reports(DB_PATH)

            if reports_df_all.empty:
                 st.info("No reports defined to delete metadata for.")
            else:
                st.write("Select reports whose metadata you want to delete:")
                selected_reports_for_deletion = []
                for index, row in reports_df_all.iterrows():
                    report_name = row["report_name"]
                    checkbox_key = f"delete_report_metadata_checkbox_{report_name}"
                    if st.checkbox(report_name, key=checkbox_key):
                        selected_reports_for_deletion.append(report_name)

                if selected_reports_for_deletion:
                    num_selected = len(selected_reports_for_deletion)
                    st.warning(f"You have selected {num_selected} report(s). Deleting metadata cannot be undone.")
                    if st.button(f"Initiate Metadata Deletion for Selected ({num_selected})", key="initiate_multi_delete_checkbox_button"):
                        st.session_state.pending_delete_reports = selected_reports_for_deletion
                        st.rerun()

            st.markdown("---")


            # --- Report Last Upload Summary ---
            st.markdown("### Report Last Upload Summary")
            with sqlite3.connect(DB_PATH) as conn:
                logs_summary = pd.read_sql_query("""
                    SELECT report_name, MAX(uploaded_at) AS last_refresh
                    FROM upload_log
                    GROUP BY report_name
                    ORDER BY last_refresh DESC
                """, conn)

            if logs_summary.empty:
                st.info("No uploads recorded for any report.")
            else:
                st.dataframe(logs_summary)

            st.markdown("---")

            # --- Report-Specific Upload Details (Single Select & Data Editor with Checkbox Workaround) ---
            st.markdown("### Report-Specific Upload Details")
            report_names_with_uploads = logs_summary["report_name"].tolist()

            if not report_names_with_uploads:
                 st.info("Upload history is empty. Upload a file first to see details.")
            else:
                 selected_report_history = st.selectbox(
                     "Select report to inspect history",
                     report_names_with_uploads,
                     key="selectbox_inspect_history"
                )

                 # Fetch detailed logs using the correct 'id' column
                 with sqlite3.connect(DB_PATH) as conn:
                     report_logs_detail = pd.read_sql_query("""
                         SELECT filename, table_alias, uploaded_at, rows, cols, id
                         FROM upload_log
                         WHERE report_name = ?
                         ORDER BY uploaded_at DESC
                     """, conn, params=(selected_report_history,))

                 st.markdown(f"#### 📁 Files uploaded for `{selected_report_history}`")
                 if report_logs_detail.empty:
                      st.info(f"No specific upload logs found for '{selected_report_history}'.")
                 else:
                      # --- WORKAROUND for selection using a Checkbox Column ---
                      # 1. Rename 'id' to 'upload_id' for display (optional, but matches previous display)
                      report_logs_detail_display = report_logs_detail.rename(columns={'id': 'upload_id'})

                      # 2. Create a copy and insert the 'Select' checkbox column at the beginning
                      df_with_selections = report_logs_detail_display.copy()
                      df_with_selections.insert(0, "Select", False) # Insert boolean column

                      st.write("Tick the box next to the upload(s) you want to delete:")

                      # 3. Use st.data_editor with column_config for the checkbox
                      # Disable editing on original columns
                      # Get the names of the original columns BEFORE inserting 'Select'
                      original_columns = report_logs_detail_display.columns.tolist()
                      disabled_columns = original_columns # Disable editing on all original columns


                      edited_df = st.data_editor(
                          df_with_selections, # Pass the dataframe *with* the Select column
                          use_container_width=True,
                          hide_index=True, # Hide pandas index
                          # Configure the Select column as a checkbox
                          # Also configure other columns like upload_id to be non-editable if desired
                          column_config={
                              "Select": st.column_config.CheckboxColumn(required=True, help="Select this upload record for deletion"),
                              "upload_id": st.column_config.NumberColumn("Upload ID", disabled=True), # Make ID column non-editable in editor
                              "filename": st.column_config.TextColumn("Filename", disabled=True),
                              "table_alias": st.column_config.TextColumn("Table Alias", disabled=True),
                              "uploaded_at": st.column_config.DatetimeColumn("Uploaded At", disabled=True),
                              "rows": st.column_config.NumberColumn("Rows", disabled=True),
                              "cols": st.column_config.NumberColumn("Cols", disabled=True),
                          },
                          disabled=disabled_columns, # Disable editing on all original data columns
                          # REMOVE selection="multiple-rows" <-- This parameter is NOT USED in this workaround
                          key=f"data_editor_upload_history_checkbox_{selected_report_history}" # Unique key
                      )

                      # 4. Filter the edited dataframe to find selected rows (where 'Select' is True)
                      # This filters the dataframe returned by the editor based on the checkbox column
                      selected_rows = edited_df[edited_df.Select == True]


                      if not selected_rows.empty:
                          # 5. Get the actual 'id' (DB primary key) from the selected rows
                          # The 'upload_id' column in the edited_df (which came from original 'id')
                          selected_upload_ids = selected_rows['upload_id'].tolist()

                          num_selected_uploads = len(selected_upload_ids)
                          st.warning(f"You have selected {num_selected_uploads} upload record(s). Initiating deletion will show a final confirmation step.")

                          # Button to initiate deletion of selected uploads
                          if st.button(f"Initiate Deletion of Selected Uploads ({num_selected_uploads})", key="initiate_delete_uploads_button"):
                              st.session_state.pending_delete_uploads = selected_upload_ids
                              st.rerun()

                      # --- End WORKAROUND ---


            st.markdown("---")

            # --- Danger Zone: Delete ALL ---
            st.markdown("### Danger Zone: Delete ALL Reports AND Data")
            st.warning("This will permanently delete ALL report definitions, ALL metadata, AND ALL data tables created by uploads!")

            confirm_delete_all_data = st.checkbox("I understand this will delete ALL data tables.", key="confirm_delete_all_data_checkbox")
            confirm_delete_all_reports = st.checkbox("I also confirm deletion of ALL report definitions and metadata.", key="confirm_delete_all_reports_checkbox")


            if confirm_delete_all_data and confirm_delete_all_reports and st.button("🔥 Execute Delete EVERYTHING", key="execute_delete_all_button"):
                st.info("Starting global deletion...")
                with sqlite3.connect(DB_PATH) as conn:
                    cursor = conn.cursor()
                    st.info("Deleting all data tables...")
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                    all_tables = [row[0] for row in cursor.fetchall()]
                    metadata_tables = ['upload_log', 'file_alias_map', 'alias_upload_status', 'sheet_rules', 'transform_rules', 'reports', 'report_structure', 'report_cutoff_log', 'sqlite_sequence', 'geometry_columns', 'spatial_ref_sys']
                    data_tables_to_drop = [tbl for tbl in all_tables if tbl not in metadata_tables]

                    dropped_count = 0
                    for table in data_tables_to_drop:
                        try:
                            cursor.execute(f"DROP TABLE IF EXISTS {table};")
                            st.write(f"Dropped table: `{table}`")
                            dropped_count += 1
                        except Exception as drop_e:
                             st.error(f"Error dropping table `{table}`: {drop_e}")

                    conn.commit()
                    st.success(f"Completed dropping {dropped_count} data table(s).")

                    st.info("Deleting all report metadata...")
                    try:
                        conn.execute("DELETE FROM reports")
                        conn.execute("DELETE FROM upload_log")
                        conn.execute("DELETE FROM report_structure")
                        conn.execute("DELETE FROM report_cutoff_log")
                        conn.execute("DELETE FROM file_alias_map")
                        conn.execute("DELETE FROM alias_upload_status")
                        conn.execute("DELETE FROM sheet_rules")
                        conn.execute("DELETE FROM transform_rules")
                        conn.commit()
                        st.success("🧨 All report metadata (definitions, logs, rules) were deleted.")
                    except Exception as e:
                        conn.rollback()
                        st.error(f"Error deleting metadata tables: {e}")
                        print(traceback.format_exc())

                st.success("Global deletion process finished.")
                st.toast("All data and metadata deleted!", icon="💥")
                st.rerun()


    except Exception as e:
        st.error(f"💥 An error occurred in the Upload History & Management section: {e}")
        print(traceback.format_exc())
        
# ──────────────────────────────────────────────────
# TAB 5: Admin – Report Structure
# selected_section == "report_structure"
# ──────────────────────────────────────────────────


elif selected_section == "report_structure":

    from ingestion.db_utils import (
        get_suggested_structure,   # ⬅️ NEW helper
        define_expected_table      # re-use save helper
    )

    st.subheader("🛠️ Admin – Define Required Tables for a Report")

    # --------------------------------------------------
    # 1) Choose or create a report
    # --------------------------------------------------
    # Fetch reports and handle report selection/creation
    reports = get_all_reports(DB_PATH)

    # Handle different return types from get_all_reports
    if reports is not None and not reports.empty:
        # If get_all_reports returns a DataFrame
        if hasattr(reports, 'to_dict'):
            # Convert DataFrame to list of dicts or extract report names
            if 'report_name' in reports.columns:
                report_names = ["-- Create new --"] + reports['report_name'].tolist()
            else:
                # If it's a DataFrame but structure is unclear, convert to list of dicts
                reports_list = reports.to_dict('records')
                report_names = ["-- Create new --"] + [report.get("report_name", str(report)) for report in reports_list]
        elif isinstance(reports, list) and len(reports) > 0:
            # If get_all_reports returns a list
            if isinstance(reports[0], dict):
                report_names = ["-- Create new --"] + [report["report_name"] for report in reports]
            else:
                report_names = ["-- Create new --"] + reports
        else:
            # Fallback for other data types
            report_names = ["-- Create new --"]
    else:
        # If reports is None or empty
        report_names = ["-- Create new --"]

     # Selectbox for choosing a report
    chosen_report = st.selectbox("Select Report", report_names, key="audit_data_report_select")

     # Handle new report creation
    if chosen_report == "-- Create new --":
        new_report_name = st.text_input("New Report Name")
        if st.button("➕ Create Report"):
            if new_report_name.strip():
                try:
                    create_new_report(new_report_name.strip(), DB_PATH)
                    st.success(f"Report '{new_report_name}' created.")
                    st.rerun()  # Rerun to refresh the report list
                except ValueError as e:
                    st.error(f"Failed to create report: {e}")
            else:
                st.warning("Please enter a report name.")
        st.stop()
    # --------------------------------------------------
    # 2) Current structure
    # --------------------------------------------------
    with sqlite3.connect(DB_PATH) as conn:
        structure_df = pd.read_sql_query(
            """
            SELECT id, table_alias, required, expected_cutoff
            FROM report_structure
            WHERE report_name = ?
            ORDER BY table_alias
            """,
            conn, params=(chosen_report,)
        )

    st.markdown("### Current Structure")
    if structure_df.empty:
        st.info("No aliases defined yet.")
    else:
        st.dataframe(structure_df, use_container_width=True)

    st.markdown("---")

    # --------------------------------------------------
    # 2b) Inline edit existing rows
    # --------------------------------------------------
    if not structure_df.empty:
        st.markdown("### ✏️ Edit existing rows")
        editable_df = structure_df.copy()

        edited_df = st.data_editor(
            editable_df,
            column_config={
                "id": st.column_config.NumberColumn(disabled=True),
                "table_alias": st.column_config.TextColumn(disabled=True),
                "required": st.column_config.CheckboxColumn(),
                "expected_cutoff": st.column_config.TextColumn(),
            },
            hide_index=True,
            use_container_width=True,
            num_rows="fixed",  # prevent adding new rows here
        )

        if st.button("💾 Save Changes to Existing Rows"):
            diff_ct = 0
            with sqlite3.connect(DB_PATH) as conn:
                cur = conn.cursor()
                for _, row in edited_df.iterrows():
                    orig = structure_df.loc[structure_df.id == row["id"]].iloc[0]
                    if (
                        int(row["required"]) != int(orig["required"])
                        or (row["expected_cutoff"] or "") != (orig["expected_cutoff"] or "")
                    ):
                        cur.execute(
                            """
                            UPDATE report_structure
                            SET required = ?, expected_cutoff = ?
                            WHERE id = ?
                            """,
                            (int(row["required"]), row["expected_cutoff"], int(row["id"])),
                        )
                        diff_ct += 1
                conn.commit()

            if diff_ct:
                st.success(f"Updated {diff_ct} row(s).")
                st.rerun()
            else:
                st.info("No changes detected.")

    # --------------------------------------------------
    # 3) Suggested aliases from upload_log (NEW)
    # --------------------------------------------------
    suggested = get_suggested_structure(chosen_report, DB_PATH)
    if suggested:
        st.markdown("### 🚀 Suggested aliases from recent uploads")
        st.write("Tick the aliases you want to add as **required**:")
        accept_states = {}
        for alias in suggested:
            accept_states[alias] = st.checkbox(f"Add `{alias}`", value=True, key=f"suggest_{alias}")

        if st.button("➕ Accept Selected Suggestions"):
            added = 0
            for alias, do_add in accept_states.items():
                if do_add:
                    define_expected_table(chosen_report, alias, required=True, db_path=DB_PATH)
                    added += 1
            st.success(f"Added {added} alias(es) to report structure.")
            st.rerun()
    else:
        st.info("No new aliases found in uploads that are missing from the structure.")

    st.markdown("---")

    # --------------------------------------------------
    # 4) Manual Add / Edit
    # --------------------------------------------------
    st.markdown("### ➕ Add or Edit an Alias")

    col1, col2, col3 = st.columns([0.4, 0.2, 0.4])
    with col1:
        alias_input = st.text_input("Table alias", key="alias_input")
    with col2:
        required_input = st.checkbox("Required?", value=True, key="required_input")
    with col3:
        cutoff_input = st.text_input("Expected cutoff (free text)", placeholder="e.g. Month-End", key="cutoff_input")

    if st.button("💾 Save Alias"):
        alias = alias_input.strip()
        if not alias:
            st.warning("Alias cannot be empty.")
        else:
            try:
                define_expected_table(
                    chosen_report,
                    alias,
                    required=required_input,
                    expected_cutoff=cutoff_input or None,
                    db_path=DB_PATH
                )
                st.success(f"Alias '{alias}' saved / updated.")
                st.rerun()
            except Exception as e:
                st.error(f"Error saving alias: {e}")

    # --------------------------------------------------
    # 5) Delete aliases
    # --------------------------------------------------
    if not structure_df.empty:
        st.markdown("---")
        st.markdown("### 🗑️ Delete Aliases")

        ids_to_del = st.multiselect(
            "Select IDs to delete",
            structure_df["id"].tolist(),
            format_func=lambda x: f"{x} – "
                + structure_df.loc[structure_df.id == x, "table_alias"].values[0],
        )

        if ids_to_del and st.button(f"Delete {len(ids_to_del)} alias(es)"):
            with sqlite3.connect(DB_PATH) as conn:
                placeholders = ",".join("?" for _ in ids_to_del)
                conn.execute(
                    f"DELETE FROM report_structure WHERE id IN ({placeholders})",
                    ids_to_del,
                )
                conn.commit()
            st.success("Selected aliases deleted.")
            st.rerun()

    # ------------------------------------------------------------------
    # 6) parameters editor
    # ------------------------------------------------------------------
    st.markdown("## ⚙️ Report parameters")
    DEFAULTS_BY_REPORT = {                       # only used if nothing in DB yet
        "Quarterly_Report": { "H2020_positions": ['all','EXPERTS','ADG','STG','POC','COG','SYG'],
            "HEU_positions":   ['all','EXPERTS','ADG','STG','POC','COG','SYG'],
            "list_H2020_payTypes": ['Interim Payments','Final Payments','Experts Payment '],
            "list_HE_payTypes":    ['Prefinancing Payments','Interim Payments','Final Payments','Experts Payment '],
            "calls_list": [
                'ERC-2023-POC (01/23)','ERC-2023-POC (04/23)','ERC-2023-POC (09/23)',
                'ERC-2023-STG','ERC-2023-COG','ERC-2023-SyG','ERC-2023-ADG',
                'ERC-2024-POC (03/24)','ERC-2024-STG','ERC-2024-COG','ERC-2024-SyG','ERC-2024-ADG','ERC-2024-POC (09/24)'
            ],
            "TTI_Targets" : {
                'ERC-2023-POC (01/23)' : 106, 'ERC-2023-POC (04/23)' : 98,'ERC-2023-POC (09/23)':97,
                'ERC-2023-STG': 304,'ERC-2023-COG' : 309,'ERC-2023-SyG' : 371,'ERC-2023-ADG' : 332,
                'ERC-2024-POC (03/24)' : 106,'ERC-2024-STG' : 300,'ERC-2024-COG' : 309,'ERC-2024-SyG' : 371, 'ERC-2024-ADG':332,
                'ERC-2024-POC (09/24)':106

            },
            "TTS_Targets" : {
                'ERC-2023-POC (01/23)' : 120, 'ERC-2023-POC (04/23)' : 120,'ERC-2023-POC (09/23)':120,
                'ERC-2023-STG': 120,'ERC-2023-COG' : 120,'ERC-2023-SyG' : 140,'ERC-2023-ADG' : 120,
                'ERC-2024-POC (03/24)' : 120,'ERC-2024-STG' : 120,'ERC-2024-COG' : 120,'ERC-2024-SyG' :140,'ERC-2024-ADG':120,
                'ERC-2024-POC (09/24)':120
            },
            "TTG_Targets" : {
                'ERC-2023-POC (01/23)' : 226, 'ERC-2023-POC (04/23)' : 218,'ERC-2023-POC (09/23)':217,
                'ERC-2023-STG': 424,'ERC-2023-COG' : 429,'ERC-2023-SyG' : 511,'ERC-2023-ADG' : 452,
                'ERC-2024-POC (03/24)' :226,'ERC-2024-STG' : 420,'ERC-2024-COG' : 429,'ERC-2024-SyG' :511, 'ERC-2024-ADG':452,
                'ERC-2024-POC (09/24)':226
            },
            "EXCLUDE_TOPICS" : ["ERC-2023-SJI-1", "ERC-2023-SJI","ERC-2024-PERA","HORIZON-ERC-2022-VICECHAIRS-IBA","HORIZON-ERC-2023-VICECHAIRS-IBA",
            "HORIZON-ERC-2025-NCPS-IBA"
           ],
           "TABLE_COLORS" :{ "BLUE" : "#2773C5" , "LIGHT_BLUE" : "#B9DBFF" , "GRID_CLR" : "#004A99" , "DARK_BLUE" :"#1B5390",  "DARK_GREY" : '#242425',
           "heading_background_color": "#1B5390","row_group_background_color": "#d6e6f4", "border_color": "#01244B", "stub_background_color": "#d6e6f4",
           "body_background_color": "#ffffff", "subtotal_background_color": "#E6E6FA", "text_color": "#01244B"
           },
           "HEU_Calls" : ['ERC-2021-ADG','ERC-2021-COG','ERC-2021-STG','ERC-2022-ADG','ERC-2022-COG' , 'ERC-2022-POC1', 'ERC-2022-POC2 (02/22)', 'ERC-2022-POC2 (05/22)','ERC-2022-POC2 (09/22)',
                        'ERC-2022-STG','ERC-2022-SYG', 'ERC-2023-ADG','ERC-2023-COG','ERC-2023-POC (01/23)','ERC-2023-POC (04/23)','ERC-2023-POC (09/23)','ERC-2023-STG','ERC-2023-SyG',
                        'ERC-2024-POC (03/24)','ERC-2024-POC (09/24)','ERC-2024-STG','ERC-2024-SyG','ERC-2024-ADG'

           ],
           "Budget_Impl" : {
               'row1_vacancy' : '1.09%' , 'row2_commit' :'100% (out of € 2.24 bln)', 'row3_L2/L1': '100 % (out of € 2.0  bln)', 'row4_paym': '100 % (out of € 2.174 bln)', 'row5_EARN': '94.21 % (out of € 0.984 bln)'
           },
           "Vacancy_Rate" : {'current_year' : '2%', "previous_year" : "1.09%"},
           "TTP_NET_HISTORY" : { 
                                 'H2020': {'IP' : 22.1, 'FP' : 48, 'Experts':41 , 'H2020': 30.8},
                                 'HEU': {'PF': 5.4, 'IP' : 15, 'FP' : 42.7, 'Experts':9.8, 'HEU': 9.9},
                                 'ALL' :{'TOTAL':15.1}
                                  },
            "TTP_GROSS_HISTORY" : { 
                                 'H2020': {'IP' : 74.2, 'FP' : 78.2, 'Experts':148 , 'H2020': 75.6},
                                 'HEU': {'PF': 5.4, 'IP' :17.7, 'FP' : 47.6, 'Experts':10, 'HEU': 10.3},
                                 'ALL' :{'TOTAL':26.7}
                                  },

            "PAYMENTS_ON_TIME_HISTORY" : { 
                                 'H2020': {'IP' : 1, 'FP' : 1, 'Experts':0 , 'H2020': 0.9996},
                                 'HEU': {'PF': 0.9933, 'IP' :1, 'FP' : 1, 'Experts':0.9943, 'HEU': 0.9943},
                                 'ALL' :{'TOTAL':0.996}
                                  },

            "Administrative_expenditure_Effectiveness" : { 
                                 'Current': 0.985,
                                 'Old': 0.975,
                                 'Target' :"99% in 30 days"
                                  },
            "Expert_meetings_Effectiveness" : { 
                                 'Current': 'na',
                                 'Old':'na',
                                 'Target' :'na'
                                  },
              "Administrative_expenditure_ttp" : { 
                                 'Current': 7.8,
                                 'Old': 7.6,
                                 'Target' :"30 days"
                                  },
            "Expert_meetings_ttp" : { 
                                 'Current': 'na',
                                 'Old':'na',
                                 'Target' : "30 days"
                                  },
                        
                
               
      }
    }
    params = load_report_params(chosen_report, DB_PATH) or DEFAULTS_BY_REPORT.get(chosen_report, {})

    # ----- defaults only for the selected report ----------------------
    DEFAULTS_FOR_THIS = DEFAULTS_BY_REPORT.get(chosen_report, {})

    # ------------------------------------------------------------------
    # 0) bootstrap the authoritative dict just once
    # ------------------------------------------------------------------
   
    if "params_full" not in st.session_state:
        # first pull what we already have in the database
        from_db = load_report_params(chosen_report, DB_PATH)  # may be {}
        # DB values win over defaults
        st.session_state.params_full = {**DEFAULTS_FOR_THIS, **from_db}

    # make a short alias for readability
    params = st.session_state.params_full


    # ---------- editable JSON textarea ----------------------------------
    txt_key = "param_editor"           # 1-line alias for the textarea’s key

    # --- if a widget (Add-key / Quick-add) set a _params_draft, promote it
    if "_params_draft" in st.session_state:
        st.session_state[txt_key] = st.session_state.pop("_params_draft")
        st.toast("✅ Parameters updated – review & save", icon="🎉")

    # Ensure the textarea has an initial value (first render only)
    # if txt_key not in st.session_state:
    #     refresh_editor_from_params()

    # Ensure initial value (BEFORE widget renders)
    if txt_key not in st.session_state:
        st.session_state[txt_key] = json.dumps(params, indent=2)

    # ---------- editable JSON textarea --------------------------------
    txt = st.text_area(
        "Edit JSON",
        key=txt_key,
        height=320,
    )

    # Always keep `params` as the object that drives the UI -------------
    try:
        params: dict = json.loads(st.session_state[txt_key] or "{}")
    except json.JSONDecodeError as e:
        st.error(f"JSON syntax error: {e}")
        st.stop()


    # -------------- Add new top-level key --------------
    with st.expander("➕ Add a new top-level key", expanded=False):
        new_key = st.text_input("Key name", key="new_top_key")
        init_type = st.selectbox("Initial type", ["list", "dict", "str", "int"], key="init_type")

        if st.button("Create key"):
            if new_key in params:
                st.warning("Key already exists!")
            else:
                params[new_key] = [] if init_type == "list" else {} if init_type == "dict" else "" if init_type == "str" else 0
                st.session_state["_params_draft"] = json.dumps(params, indent=2)
                st.session_state.pop("new_top_key", None)
                st.toast("✅ Key created")
                st.rerun()

    st.markdown("---")

    # ---------- 4. QUICK PATCH EXISTING KEYS  -----------------------------------
    # -------------------------------------------------------------------
    # 4. QUICK PATCH / EDIT / DELETE
    # -------------------------------------------------------------------
    st.markdown("### 🔄 Quick add / update")

    # def commit_and_rerun(params_dict: dict, toast_msg: str):
    #     """
    #     Write json draft → session-state, toast, force rerun
    #     (we cannot touch param_editor while it’s mounted)
    #     """
    #     st.session_state["_params_draft"] = json.dumps(params_dict, indent=2)
    #     st.toast(toast_msg, icon="✅")
    #     st.rerun()
    def commit_and_rerun(params_dict: dict, toast_msg: str):
        """
        Save each parameter to DB, update session, show toast, rerun app
        """
        for k, v in params_dict.items():
            upsert_report_param(chosen_report, k, v)

        st.session_state["_params_draft"] = json.dumps(params_dict, indent=2)
        st.toast(toast_msg, icon="✅")
        st.rerun()

    gkey = st.selectbox("Existing key", list(params.keys()), key="qa_sel")

    # detect the value type of the selected key
    current_val = params[gkey]
    is_list = isinstance(current_val, list)
    is_dict = isinstance(current_val, dict)

    vtype = st.selectbox(
        "Action",
        [
            "replace scalar",            # for str / number
            "append to list",            # for list
            "edit list item",            # for list
            "delete list item",          # for list
            "add / update dict entry",   # for dict
            "delete dict entry"          # for dict
        ],
        key="qa_action"
    )

    # ───────── scalar replace ──────────────────────────────────────────
    if vtype == "replace scalar":
        if is_list or is_dict:
            st.warning("Selected key is not a scalar.")
        else:
            new_val = st.text_input("New value", key="qa_scalar_in")
            if st.button("💾 Replace value", key="qa_scalar_btn"):
                params[gkey] = new_val
                commit_and_rerun(params, "Scalar replaced")

    # ───────── list append  ────────────────────────────────────────────
    elif vtype == "append to list":
        if not is_list:
            st.warning("Selected key is not a list.")
        else:
            new_item = st.text_input("Item to append", key="qa_append_in")
            if st.button("➕ Append", key="qa_append_btn"):
                params[gkey].append(new_item)
                commit_and_rerun(params, "Item appended")

    # ───────── list edit  ──────────────────────────────────────────────
    elif vtype == "edit list item":
        if not is_list:
            st.warning("Selected key is not a list.")
        else:
            sel = st.selectbox("Pick item", current_val, key="qa_list_pick")
            new_item = st.text_input("New value", value=sel, key="qa_list_edit")
            if st.button("✏️ Update item", key="qa_list_edit_btn"):
                idx = current_val.index(sel)
                params[gkey][idx] = new_item
                commit_and_rerun(params, "List item updated")

    # ───────── list delete  ────────────────────────────────────────────
    elif vtype == "delete list item":
        if not is_list:
            st.warning("Selected key is not a list.")
        else:
            sel = st.selectbox("Pick item to delete", current_val, key="qa_list_del_pick")
            if st.button("🗑️ Delete item", key="qa_list_del_btn"):
                params[gkey].remove(sel)
                commit_and_rerun(params, "List item deleted")

    # ───────── dict upsert  ────────────────────────────────────────────
    elif vtype == "add / update dict entry":
        if not is_dict:
            st.warning("Selected key is not a dict.")
        else:
            subk = st.text_input("Sub-key", key="qa_dict_key")
            subv = st.text_input("Sub-value", key="qa_dict_val")
            if st.button("💾 Add / update", key="qa_dict_upd_btn"):
                params[gkey][subk] = subv
                commit_and_rerun(params, "Dict entry saved")

    # ───────── dict delete  ────────────────────────────────────────────
    elif vtype == "delete dict entry":
        if not is_dict:
            st.warning("Selected key is not a dict.")
        else:
            subk = st.selectbox("Pick sub-key to delete", list(current_val.keys()), key="qa_dict_del_pick")
            if st.button("🗑️ Delete entry", key="qa_dict_del_btn"):
                params[gkey].pop(subk, None)
                commit_and_rerun(params, "Dict entry deleted")


    st.markdown("---")

    # ---------- 5. CALL-TARGETS MANAGER (optional) -------------------------------
    if isinstance(params.get("call_targets"), dict):
        st.markdown("### 🎯 Maintain `call_targets`")

        ct_dict = params["call_targets"]
        call_names = list(ct_dict.keys())

        sel = st.selectbox("Select call (or '-- New call --')",
                        ["-- New call --"] + call_names, key="ct_sel")

        call_key = st.text_input("Call name", value="" if sel == "-- New call --" else sel,
                                key="ct_name")
        if call_key:
            entry = ct_dict.get(call_key, {"TTI": "", "TTS": "", "TTG": ""})
            c1,c2,c3 = st.columns(3)
            with c1:  tti = st.text_input("TTI", value=str(entry["TTI"]), key="ct_tti")
            with c2:  tts = st.text_input("TTS", value=str(entry["TTS"]), key="ct_tts")
            with c3:  ttg = st.text_input("TTG", value=str(entry["TTG"]), key="ct_ttg")

            def _num(v):
                try: return int(v) if float(v).is_integer() else float(v)
                except: return v.strip()

            if st.button("💾 Save / update call", key="ct_save"):
                ct_dict[call_key] = {"TTI": _num(tti), "TTS": _num(tts), "TTG": _num(ttg)}
                params["call_targets"] = ct_dict
                st.session_state[txt_key] = json.dumps(params, indent=2)
                st.success("Saved (remember to **Save parameters to DB**)")

            if sel != "-- New call --" and st.button("🗑️ Delete call", key="ct_del"):
                del ct_dict[call_key]
                params["call_targets"] = ct_dict
                st.session_state[txt_key] = json.dumps(params, indent=2)
                st.success("Deleted (remember to **Save parameters to DB**)")

        st.markdown("---")

 
    # --------------------------------------------------
    # 7) ☰ Report-Modules mapping
    # --------------------------------------------------
    from ingestion.db_utils import (
        list_report_modules,
        upsert_report_module,
        delete_report_module,
        ensure_report_modules_table
    )
    import sqlite3
    import logging
    import streamlit as st
    import pandas as pd
    import importlib

    # Set up logging
    logging.basicConfig(level=logging.DEBUG)
    logger = logging.getLogger(__name__)

    # Ensure the report_modules table exists
    logger.debug(f"Ensuring report_modules table exists in {DB_PATH}")
    ensure_report_modules_table(DB_PATH)

    # Verify the table exists
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='report_modules';")
            table_exists = cursor.fetchone()
            if table_exists:
                logger.debug("report_modules table exists.")
            else:
                logger.error("report_modules table does not exist after ensure_report_modules_table.")
                st.error("⚠️ Database error: report_modules table could not be created.")
                st.stop()
    except sqlite3.Error as e:
        logger.error(f"Database error while verifying table existence: {str(e)}")
        st.error(f"⚠️ Database error: {str(e)}")
        st.stop()

    # Define report-to-module mapping
    report_to_module = {
        "Quarterly_Report": "reporting.quarterly_report",
        # Add more reports here as needed, e.g.,
        # "Annual_Report": "reporting.annual_report",
        # "Invoice_Summary": "reporting.invoice_summary",
    }

    # Load modules for each report
    available_modules = {}
    for report_name, mod_path in report_to_module.items():
        try:
            mod = importlib.import_module(mod_path)
            MODULES = getattr(mod, "MODULES", {})
            if not MODULES:
                logger.error(f"No MODULES found in {mod_path} for {report_name}.")
                st.error(f"⚠️ No modules found for report '{report_name}'. Please check the module registry.")
                st.stop()
            available_modules[report_name] = MODULES
            logger.debug(f"Loaded MODULES for {report_name}: {list(MODULES.keys())}")
        except ImportError as e:
            logger.error(f"Error loading module {mod_path} for {report_name}: {e}")
            st.error(f"⚠️ Failed to load modules for report '{report_name}': {e}")
            st.stop()
        except Exception as e:
            logger.error(f"Unexpected error loading module {mod_path} for {report_name}: {e}")
            st.error(f"⚠️ Unexpected error loading modules for report '{report_name}': {e}")
            st.stop()

    # Function to fetch all mappings from report_modules table
    def fetch_all_report_modules(db_path):
        try:
            with sqlite3.connect(db_path) as conn:
                query = "SELECT * FROM report_modules ORDER BY report_name, run_order"
                df = pd.read_sql_query(query, conn)
                logger.debug(f"Fetched {len(df)} report modules from database")
                return df
        except sqlite3.Error as e:
            logger.error(f"Error fetching report modules: {str(e)}")
            st.error(f"⚠️ Database error: {str(e)}")
            return pd.DataFrame()

    # Force refresh data if we just performed a delete operation
    if "force_refresh_mappings" in st.session_state:
        st.session_state.pop("force_refresh_mappings")
        st.cache_data.clear()  # Clear any cached data

    # Enhanced delete function with aggressive debugging
    def delete_mapping_debug(mapping_id, db_path):
        """Debug version with extensive logging"""
        try:
            logger.debug(f"=== STARTING DELETE OPERATION FOR ID: {mapping_id} ===")
            
            # Step 1: Verify mapping exists
            with sqlite3.connect(db_path) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM report_modules WHERE id = ?", (mapping_id,))
                existing_mapping = cursor.fetchone()
                
                if not existing_mapping:
                    return False, f"No mapping found with ID: {mapping_id}"
                
                logger.debug(f"Found mapping: {existing_mapping}")
            
            # Step 2: Try direct SQL delete (bypass the original function)
            logger.debug("Attempting direct SQL delete...")
            with sqlite3.connect(db_path) as conn:
                cursor = conn.cursor()
                
                # Get count before
                cursor.execute("SELECT COUNT(*) FROM report_modules WHERE id = ?", (mapping_id,))
                count_before = cursor.fetchone()[0]
                logger.debug(f"Count before delete: {count_before}")
                
                # Perform delete
                cursor.execute("DELETE FROM report_modules WHERE id = ?", (mapping_id,))
                rows_affected = cursor.rowcount
                logger.debug(f"Rows affected by DELETE: {rows_affected}")
                
                # IMPORTANT: Explicit commit
                conn.commit()
                logger.debug("Explicit COMMIT executed")
                
                # Get count after
                cursor.execute("SELECT COUNT(*) FROM report_modules WHERE id = ?", (mapping_id,))
                count_after = cursor.fetchone()[0]
                logger.debug(f"Count after delete: {count_after}")
                
                # Get all remaining IDs to see what's in the table
                cursor.execute("SELECT id FROM report_modules ORDER BY id")
                all_ids = [row[0] for row in cursor.fetchall()]
                logger.debug(f"All remaining IDs: {all_ids}")
                
                if count_after == 0 and rows_affected > 0:
                    return True, f"Direct SQL delete successful (ID: {mapping_id})"
                else:
                    return False, f"Direct SQL delete failed - before: {count_before}, after: {count_after}, affected: {rows_affected}"
                    
        except sqlite3.Error as e:
            logger.error(f"Database error: {str(e)}")
            return False, f"Database error: {str(e)}"
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return False, f"Unexpected error: {str(e)}"

    st.markdown("## ⚙️ Manage Report-Module Mappings")

    # Add database debugging section
    with st.expander("🔍 Database Debug Tools", expanded=False):
        st.markdown("### Database Integrity Check")
        
        if st.button("🔍 Check Database State"):
            try:
                with sqlite3.connect(DB_PATH) as conn:
                    cursor = conn.cursor()
                    
                    # Check if table exists
                    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='report_modules'")
                    table_exists = cursor.fetchone()
                    st.write(f"**Table exists:** {bool(table_exists)}")
                    
                    if table_exists:
                        # Get table schema
                        cursor.execute("PRAGMA table_info(report_modules)")
                        schema = cursor.fetchall()
                        st.write("**Table schema:**")
                        st.write(schema)
                        
                        # Get all records
                        cursor.execute("SELECT * FROM report_modules ORDER BY id")
                        all_records = cursor.fetchall()
                        st.write(f"**Total records:** {len(all_records)}")
                        
                        if all_records:
                            df_debug = pd.DataFrame(all_records, columns=[col[1] for col in schema])
                            st.dataframe(df_debug)
                        
                        # Check for any constraints or triggers
                        cursor.execute("SELECT sql FROM sqlite_master WHERE type='trigger' AND tbl_name='report_modules'")
                        triggers = cursor.fetchall()
                        if triggers:
                            st.write("**Triggers found:**")
                            for trigger in triggers:
                                st.code(trigger[0])
                        else:
                            st.write("**No triggers found**")
                            
            except Exception as e:
                st.error(f"Database check failed: {e}")
        
        if st.button("🧪 Test Direct Delete"):
            try:
                with sqlite3.connect(DB_PATH) as conn:
                    cursor = conn.cursor()
                    
                    # Create a test record
                    cursor.execute("""
                        INSERT INTO report_modules (report_name, module_name, run_order, enabled) 
                        VALUES ('TEST_REPORT', 'TEST_MODULE', 999, 1)
                    """)
                    test_id = cursor.lastrowid
                    conn.commit()
                    st.write(f"Created test record with ID: {test_id}")
                    
                    # Try to delete it
                    cursor.execute("DELETE FROM report_modules WHERE id = ?", (test_id,))
                    rows_affected = cursor.rowcount
                    conn.commit()
                    
                    # Check if it's gone
                    cursor.execute("SELECT COUNT(*) FROM report_modules WHERE id = ?", (test_id,))
                    still_exists = cursor.fetchone()[0]
                    
                    st.write(f"**Delete test results:**")
                    st.write(f"- Rows affected: {rows_affected}")
                    st.write(f"- Record still exists: {still_exists > 0}")
                    
                    if still_exists == 0:
                        st.success("✅ Direct delete works fine!")
                    else:
                        st.error("❌ Direct delete failed!")
                        
            except Exception as e:
                st.error(f"Test delete failed: {e}")

    # ── Display Current Mappings ────────────────────────────────────────────────
    st.markdown("### Current Mappings")
    all_mappings_df = fetch_all_report_modules(DB_PATH)
    if all_mappings_df.empty:
        st.info("No report-module mappings exist yet.")
    else:
        st.dataframe(all_mappings_df, hide_index=True, use_container_width=True)

    # ── Add New Mapping ─────────────────────────────────────────────────
    with st.expander("➕ Add New Mapping", expanded=True):
        st.markdown("### Add a New Report-Module Mapping")
        selected_report = st.selectbox("Select Report", options=list(report_to_module.keys()), key="add_report_select")
        if selected_report in available_modules and available_modules[selected_report]:
            module_options = list(available_modules[selected_report].keys())
            selected_module = st.selectbox("Select Module", options=module_options, key="add_module_select")
            run_order = st.number_input("Run Order (1 = first)", min_value=1, value=1, step=1, key="add_run_order")
            # Add unique key for the checkbox with timestamp-like suffix
            import time
            enabled = st.checkbox("Enabled", value=True, key=f"add_enabled_main")

            if st.button("💾 Add Mapping", key="add_mapping_btn"):
                try:
                    upsert_report_module(
                        report_name=selected_report,
                        module_name=selected_module,
                        run_order=run_order,
                        enabled=enabled,
                        db_path=DB_PATH
                    )
                    st.success(f"Mapping for {selected_module} added to {selected_report}.")
                    st.rerun()
                except Exception as e:
                    logger.error(f"Error adding mapping: {str(e)}")
                    st.error(f"⚠️ Failed to add mapping: {str(e)}")
        else:
            st.warning(f"No modules available for report '{selected_report}'.")

    # ── Edit Existing Mapping ─────────────────────────────────────────────
    if not all_mappings_df.empty:
        with st.expander("✏️ Edit Existing Mapping", expanded=False):
            st.markdown("### Edit an Existing Mapping")
            mapping_to_edit = st.selectbox(
                "Select Mapping to Edit", 
                options=all_mappings_df.index, 
                format_func=lambda x: f"{all_mappings_df.loc[x, 'report_name']} - {all_mappings_df.loc[x, 'module_name']} (ID: {all_mappings_df.loc[x, 'id']})",
                key="edit_mapping_selectbox"
            )
            selected_mapping = all_mappings_df.loc[mapping_to_edit]
            
            edit_report_name = selected_mapping['report_name']
            edit_module_name = st.selectbox("Module", options=list(available_modules[edit_report_name].keys()), index=list(available_modules[edit_report_name].keys()).index(selected_mapping['module_name']), key="edit_module_select")
            edit_run_order = st.number_input("Run Order", min_value=1, value=int(selected_mapping['run_order']), step=1, key="edit_run_order")
            # Add unique key for the checkbox using the mapping ID
            edit_enabled = st.checkbox("Enabled", value=bool(selected_mapping['enabled']), key=f"edit_enabled_main_{selected_mapping['id']}")

            if st.button("💾 Update Mapping", key="update_mapping_btn"):
                try:
                    upsert_report_module(
                        report_name=edit_report_name,
                        module_name=edit_module_name,
                        run_order=edit_run_order,
                        enabled=edit_enabled,
                        db_path=DB_PATH
                    )
                    st.success(f"Mapping for {edit_module_name} in {edit_report_name} updated.")
                    st.rerun()
                except Exception as e:
                    logger.error(f"Error updating mapping: {str(e)}")
                    st.error(f"⚠️ Failed to update mapping: {str(e)}")

    # ── Delete Mapping ────────────────────────────────────────────────
    if not all_mappings_df.empty:
        with st.expander("🗑️ Delete Mapping", expanded=False):
            st.markdown("### Delete a Mapping")
            
            # Get fresh data
            with sqlite3.connect(DB_PATH) as conn:
                fresh_df = pd.read_sql_query(
                    "SELECT * FROM report_modules ORDER BY report_name, run_order", 
                    conn
                )
            
            if fresh_df.empty:
                st.info("No mappings to delete.")
            else:
                # Create selection options
                options = []
                for _, row in fresh_df.iterrows():
                    option = f"{row['report_name']} - {row['module_name']} (ID: {row['id']})"
                    options.append((row['id'], option))
                
                selected = st.selectbox(
                    "Select Mapping to Delete",
                    options=options,
                    format_func=lambda x: x[1]
                )
                
                if selected and st.button("🗑️ Delete Selected", type="primary"):
                    delete_id = selected[0]
                    
                    try:
                        # Direct SQL approach
                        conn = sqlite3.connect(DB_PATH)
                        conn.execute("PRAGMA foreign_keys = OFF")  # Temporarily disable FK constraints
                        
                        cursor = conn.cursor()
                        cursor.execute("DELETE FROM report_modules WHERE id = ?", (delete_id,))
                        
                        if cursor.rowcount > 0:
                            conn.commit()
                            conn.close()
                            
                            st.success(f"✅ Deleted mapping ID {delete_id}")
                            time.sleep(1)
                            st.rerun()
                        else:
                            conn.close()
                            st.error("❌ No rows were deleted")
                            
                    except Exception as e:
                        if conn:
                            conn.close()
                        st.error(f"❌ Error: {str(e)}")
# ---------------------------------------------------
# ---------------------------------------------------
# Template Management
# ---------------------------------------------------
elif selected_section == "template_editor":
    st.subheader("🛠️ Template Management")

    # ─────────────────────── Imports ───────────────────────
    import base64
    import docx
    import docxedit
    from docx import Document
    from pathlib import Path
    import mammoth

    # ─────────────────────── Streamlit Imports ───────────────────────
    import streamlit as st

    # ─────────────────────── Constants ───────────────────────
    TEMPLATE_DIRECTORY = Path("reporting/templates/docx")
    TEMPLATE_DIRECTORY.mkdir(parents=True, exist_ok=True)
    KEY_TEMPLATE_PATH = "template_docx_path"

    # ─────────────────────── Session State Initialization ───────────────────────
    if KEY_TEMPLATE_PATH not in st.session_state:
        st.session_state[KEY_TEMPLATE_PATH] = None

    # ─────────────────────── Helper Functions ───────────────────────
    def list_template_files():
        """Return sorted list of *.docx files in the template directory."""
        return sorted(path.name for path in TEMPLATE_DIRECTORY.glob("*.docx"))

    def extract_docx_paragraphs(doc):
        """Extract paragraphs from the DOCX for display in the placeholder selector."""
        paragraphs = []
        for para in doc.paragraphs:
            text = "".join(run.text for run in para.runs)
            if text.strip():  # Only include non-empty paragraphs
                paragraphs.append(text)
        return paragraphs

    def add_placeholder_to_docx(doc, placeholder_name, para_index=None):
        """Add a Jinja2 placeholder to the DOCX at a specific paragraph using docxedit."""
        placeholder = f"{{{{ {placeholder_name} }}}}"
        if para_index is not None and 0 <= para_index < len(doc.paragraphs):
            # Use docxedit to insert the placeholder at the specified paragraph
            current_text = "".join(run.text for run in doc.paragraphs[para_index].runs)
            if current_text.strip():
                # Append the placeholder to the existing text in the paragraph
                new_text = current_text + " " + placeholder
                docxedit.replace_string(doc, old_string=current_text, new_string=new_text)
            else:
                # If the paragraph is empty, just set the placeholder
                docxedit.replace_string(doc, old_string="", new_string=placeholder)
        else:
            # Append to the end if no valid index is provided
            doc.add_paragraph(placeholder)
        return doc

    def docx_to_html(docx_path):
        """Convert DOCX to HTML using mammoth for preview."""
        try:
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value  # The converted HTML
                messages = result.messages  # Any conversion messages/warnings
                if messages:
                    st.warning("Conversion warnings: " + "; ".join(str(msg) for msg in messages))
                return html_content
        except Exception as e:
            st.error(f"Failed to convert DOCX to HTML: {e}")
            return None

    # ─────────────────────── Template Management ───────────────────────
    st.markdown("### 📄 Template File")
    template_files = ["-- new --"] + list_template_files()
    template_selector = st.selectbox(
        "Choose a file",
        template_files,
        key="template_file_selector"
    )

    if template_selector == "-- new --":
        new_file_name = st.text_input(
            "New file name",
            value="template.docx",
            help="Must end with .docx"
        )
        active_template_path = TEMPLATE_DIRECTORY / new_file_name
        # Create a blank DOCX if it doesn't exist
        if not active_template_path.exists():
            doc = Document()
            doc.add_paragraph("New Template - Edit in Word and re-upload")
            doc.save(active_template_path)
    else:
        active_template_path = TEMPLATE_DIRECTORY / template_selector

    # Load the template
    st.session_state[KEY_TEMPLATE_PATH] = str(active_template_path)

    # Upload DOCX
    uploaded_docx = st.file_uploader("Upload DOCX", type="docx", key="docx_upload")
    if uploaded_docx:
        active_template_path.write_bytes(uploaded_docx.read())
        st.session_state[KEY_TEMPLATE_PATH] = str(active_template_path)

    # ─────────────────────── Download for Editing ───────────────────────
    st.markdown("### ✏️ Edit Template")
    st.info(
        "Streamlit cannot edit DOCX files directly with pixel-perfect fidelity. "
        "Download the template, edit it in Microsoft Word (or a compatible editor), "
        "and re-upload it to make formatting changes. Use the section below to add Jinja2 placeholders."
    )
    if active_template_path.exists():
        with open(active_template_path, "rb") as f:
            docx_bytes = f.read()
        b64 = base64.b64encode(docx_bytes).decode()
        st.markdown(
            f'<a download="{active_template_path.name}" href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}">Download Template for Editing</a>',
            unsafe_allow_html=True
        )
    else:
        st.warning("No template selected or uploaded.")

    # ─────────────────────── File Actions ───────────────────────
    st.markdown("### 💾 File Actions")
    save_button, download_button, delete_button = st.columns(3)

    if save_button.button("💾 Save", disabled=not active_template_path.name.endswith(".docx")):
        st.success(f"Template is saved at: {active_template_path}")

    if download_button.button("📥 Download", disabled=not active_template_path.exists()):
        try:
            with open(active_template_path, "rb") as f:
                docx_bytes = f.read()
            b64 = base64.b64encode(docx_bytes).decode()
            st.markdown(
                f'<a download="{active_template_path.name}" href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}">Click to download</a>',
                unsafe_allow_html=True
            )
        except Exception as e:
            st.error(f"Download failed: {e}")

    if delete_button.button("🗑️ Delete", disabled=not active_template_path.exists()):
        try:
            active_template_path.unlink()
            st.warning(f"Deleted: {active_template_path.name}")
            st.session_state[KEY_TEMPLATE_PATH] = None
            st.rerun()
        except Exception as e:
            st.error(f"Delete failed: {e}")

# ---------------------------------------------------
# AUDIT INPUT
# ---------------------------------------------------
elif selected_section == "audit_data_input":
    from PIL import Image
    from docxtpl import DocxTemplate, InlineImage
    st.header("📊 Audit Data Input")
    st.write("Enter the data for External Audits and Error Rates tables below.")
    
  
    # Enhanced CSS for better form layout
    st.markdown(
        """
        <style>
        .stForm {
            max-width: 1200px;
            margin: 0 auto;
            padding: 25px;
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            border-radius: 15px;
            border: 2px solid #dee2e6;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        .stColumn {
            padding: 0 8px;
            display: flex;
            flex-direction: column;
            justify-content: flex-end;
        }
        .stTextInput > div > div > input {
            width: 100% !important;
            height: 38px !important;
            font-size: 13px !important;
            text-align: center !important;
            border: 2px solid #ced4da !important;
            border-radius: 6px !important;
            transition: border-color 0.3s ease !important;
        }
        .stTextInput > div > div > input:focus {
            border-color: #1f4788 !important;
            box-shadow: 0 0 0 2px rgba(31, 71, 136, 0.2) !important;
        }
        .column-header {
            background: linear-gradient(135deg, #1f4788 0%, #2c5aa0 100%);
            color: white;
            padding: 12px 8px;
            margin-bottom: 12px;
            border-radius: 8px;
            font-weight: bold;
            font-size: 12px;
            text-align: center;
            box-shadow: 0 2px 4px rgba(31, 71, 136, 0.3);
            border: 1px solid #1a3d6b;
        }
        .section-label {
            background: linear-gradient(135deg, #e9ecef 0%, #f8f9fa 100%);
            padding: 15px 20px;
            margin: 15px 0;
            border-radius: 10px;
            font-weight: bold;
            border-left: 5px solid #1f4788;
            border: 1px solid #dee2e6;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
            font-size: 14px;
            color: #495057;
        }
        .form-title {
            background: linear-gradient(135deg, #1f4788 0%, #2c5aa0 100%);
            color: white;
            padding: 20px;
            margin: -25px -25px 20px -25px;
            border-radius: 15px 15px 0 0;
            text-align: center;
            font-size: 18px;
            font-weight: bold;
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # Fetch reports and handle report selection/creation
    reports_df = get_all_reports(DB_PATH)
    report_names = ["-- Create new --"] + reports_df["report_name"].tolist() if not reports_df.empty else ["-- Create new --"]
    chosen_report = st.selectbox("Select Report", report_names, key="audit_data_report_select")

    # Handle new report creation
    if chosen_report == "-- Create new --":
        new_report_name = st.text_input("New Report Name")
        if st.button("➕ Create Report"):
            if new_report_name.strip():
                try:
                    create_new_report(new_report_name.strip(), DB_PATH)
                    st.success(f"Report '{new_report_name}' created.")
                    st.rerun()
                except ValueError as e:
                    st.error(f"Failed to create report: {e}")
            else:
                st.warning("Please enter a report name.")
        st.stop()
    
    # Load report parameters
    try:
        report_params = load_report_params(chosen_report, DB_PATH)
        current_year = report_params.get("current_year", datetime.now().year)
        last_date = report_params.get("last_date", datetime.now().date())
    except Exception as e:
        st.error(f"Failed to load report parameters: {e}")
        st.stop()

    def load_existing_data_or_defaults(chosen_report, current_year):
        """Load existing data from database or set default values."""
        defaults = {
            # External Audits defaults
            'target_cas': "150", 'target_joint': "N/A", 'target_subtotal': "150", 'target_court': "N/A", 'target_total': "150",
            'cumulative_cas': "1020 (900) ***", 'cumulative_joint': "N/A", 'cumulative_subtotal': "1020 (900) ***", 'cumulative_court': "N/A", 'cumulative_total': "1020 (900) ***",
            'planned_cas': "150", 'planned_joint': "0", 'planned_subtotal': "150", 'planned_court': "0", 'planned_total': "150",
            f'ongoing_{current_year}_cas': "70", f'ongoing_{current_year}_joint': "0", f'ongoing_{current_year}_subtotal': "70", f'ongoing_{current_year}_court': "5", f'ongoing_{current_year}_total': "75",
            'ongoing_prev_cas': "73", 'ongoing_prev_joint': "0", 'ongoing_prev_subtotal': "73", 'ongoing_prev_court': "0", 'ongoing_prev_total': "73",
            f'total_ongoing_{current_year}_cas': "143", f'total_ongoing_{current_year}_joint': "0", f'total_ongoing_{current_year}_subtotal': "143", f'total_ongoing_{current_year}_court': "5", f'total_ongoing_{current_year}_total': "148",
            'closed_prev_cas': "823", 'closed_prev_joint': "13", 'closed_prev_subtotal': "836", 'closed_prev_court': "61", 'closed_prev_total': "897",
            f'audited_{current_year}_cas': "8", f'audited_{current_year}_joint': "3", f'audited_{current_year}_subtotal': "11", f'audited_{current_year}_court': "15", f'audited_{current_year}_total': "26",
            f'closed_{current_year}_prev_cas': "132", f'closed_{current_year}_prev_joint': "0", f'closed_{current_year}_prev_subtotal': "132", f'closed_{current_year}_prev_court': "1", f'closed_{current_year}_prev_total': "133",
            f'total_closed_{current_year}_cas': "140", f'total_closed_{current_year}_joint': "3", f'total_closed_{current_year}_subtotal': "143", f'total_closed_{current_year}_court': "16", f'total_closed_{current_year}_total': "159",
            'total_cumulative_closed_cas': "963", 'total_cumulative_closed_joint': "16", 'total_cumulative_closed_subtotal': "979", 'total_cumulative_closed_court': "77", 'total_cumulative_closed_total': "1056",
            'total_audited_cas': "1106", 'total_audited_joint': "16", 'total_audited_subtotal': "1122", 'total_audited_court': "82", 'total_audited_total': "1204",
            # Error Rates defaults
            'cas_error_rate': "3.55", 'cas_comments': "Common Representative Error rate computed by the Common Audit Service (CAS) with top ups included. (source: SAR-Wiki)", 'cas_to_be_reported': "Quarterly basis",
            'ercea_residual_error_rate': "0.92", 'ercea_residual_comments': "ERCEA Residual error rate based on the CRS 1, 2, 3 & 4 (source: SAR-Wiki)", 'ercea_residual_to_be_reported': "Quarterly basis",
            'ercea_overall_error_rate': "1.30", 'ercea_overall_comments': "All ERCEA participations audited (source: SAR-Wiki)", 'ercea_overall_to_be_reported': "Quarterly basis"
        }
        
        # Try to load existing data
        try:
            existing_data = fetch_vars_for_report(chosen_report, DB_PATH)
            
            if existing_data:
                # Check if we have external_audits data and extract field values
                if 'external_audits' in existing_data:
                    st.success("✅ Found existing External Audits data - populating form!")
                    external_data = existing_data['external_audits']
                    
                    # Extract individual field values from the saved DataFrame structure
                    if isinstance(external_data, dict) and 'Status' in external_data:
                        # Get the data columns
                        cas_values = external_data.get('CAS', {})
                        joint_values = external_data.get('Joint with Court of auditors*', {})
                        subtotal_values = external_data.get('Subtotal for error rates and coverage', {})
                        court_values = external_data.get('Court of auditors only', {})
                        total_values = external_data.get('Total', {})
                        
                        # Map back to form fields by row index
                        # Row 0: ERCEA TARGETS
                        if '0' in cas_values:
                            defaults['target_cas'] = str(cas_values['0'])
                            defaults['target_joint'] = str(joint_values.get('0', 'N/A'))
                            defaults['target_subtotal'] = str(subtotal_values.get('0', '150'))
                            defaults['target_court'] = str(court_values.get('0', 'N/A'))
                            defaults['target_total'] = str(total_values.get('0', '150'))
                        
                        # Row 1: ERCEA TARGETS Cumulative
                        if '1' in cas_values:
                            defaults['cumulative_cas'] = str(cas_values['1'])
                            defaults['cumulative_joint'] = str(joint_values.get('1', 'N/A'))
                            defaults['cumulative_subtotal'] = str(subtotal_values.get('1', '1020 (900) ***'))
                            defaults['cumulative_court'] = str(court_values.get('1', 'N/A'))
                            defaults['cumulative_total'] = str(total_values.get('1', '1020 (900) ***'))
                        
                        # Row 2: Planned
                        if '2' in cas_values:
                            defaults['planned_cas'] = str(cas_values['2'])
                            defaults['planned_joint'] = str(joint_values.get('2', '0'))
                            defaults['planned_subtotal'] = str(subtotal_values.get('2', '150'))
                            defaults['planned_court'] = str(court_values.get('2', '0'))
                            defaults['planned_total'] = str(total_values.get('2', '150'))
                        
                        # Row 3: On-going [Launched in current year]
                        if '3' in cas_values:
                            defaults[f'ongoing_{current_year}_cas'] = str(cas_values['3'])
                            defaults[f'ongoing_{current_year}_joint'] = str(joint_values.get('3', '0'))
                            defaults[f'ongoing_{current_year}_subtotal'] = str(subtotal_values.get('3', '70'))
                            defaults[f'ongoing_{current_year}_court'] = str(court_values.get('3', '5'))
                            defaults[f'ongoing_{current_year}_total'] = str(total_values.get('3', '75'))
                        
                        # Row 4: On-going [Launched in previous years]
                        if '4' in cas_values:
                            defaults['ongoing_prev_cas'] = str(cas_values['4'])
                            defaults['ongoing_prev_joint'] = str(joint_values.get('4', '0'))
                            defaults['ongoing_prev_subtotal'] = str(subtotal_values.get('4', '73'))
                            defaults['ongoing_prev_court'] = str(court_values.get('4', '0'))
                            defaults['ongoing_prev_total'] = str(total_values.get('4', '73'))
                        
                        # Row 5: TOTAL On-going as of 31 December
                        if '5' in cas_values:
                            defaults[f'total_ongoing_{current_year}_cas'] = str(cas_values['5'])
                            defaults[f'total_ongoing_{current_year}_joint'] = str(joint_values.get('5', '0'))
                            defaults[f'total_ongoing_{current_year}_subtotal'] = str(subtotal_values.get('5', '143'))
                            defaults[f'total_ongoing_{current_year}_court'] = str(court_values.get('5', '5'))
                            defaults[f'total_ongoing_{current_year}_total'] = str(total_values.get('5', '148'))
                        
                        # Row 6: Closed in previous years
                        if '6' in cas_values:
                            defaults['closed_prev_cas'] = str(cas_values['6'])
                            defaults['closed_prev_joint'] = str(joint_values.get('6', '13'))
                            defaults['closed_prev_subtotal'] = str(subtotal_values.get('6', '836'))
                            defaults['closed_prev_court'] = str(court_values.get('6', '61'))
                            defaults['closed_prev_total'] = str(total_values.get('6', '897'))
                        
                        # Row 7: Closed in current year from audited participations launched in current year
                        if '7' in cas_values:
                            defaults[f'audited_{current_year}_cas'] = str(cas_values['7'])
                            defaults[f'audited_{current_year}_joint'] = str(joint_values.get('7', '3'))
                            defaults[f'audited_{current_year}_subtotal'] = str(subtotal_values.get('7', '11'))
                            defaults[f'audited_{current_year}_court'] = str(court_values.get('7', '15'))
                            defaults[f'audited_{current_year}_total'] = str(total_values.get('7', '26'))
                        
                        # Row 8: Closed in current year from audited participations launched in previous years
                        if '8' in cas_values:
                            defaults[f'closed_{current_year}_prev_cas'] = str(cas_values['8'])
                            defaults[f'closed_{current_year}_prev_joint'] = str(joint_values.get('8', '0'))
                            defaults[f'closed_{current_year}_prev_subtotal'] = str(subtotal_values.get('8', '132'))
                            defaults[f'closed_{current_year}_prev_court'] = str(court_values.get('8', '1'))
                            defaults[f'closed_{current_year}_prev_total'] = str(total_values.get('8', '133'))
                        
                        # Row 9: TOTAL Closed in current year
                        if '9' in cas_values:
                            defaults[f'total_closed_{current_year}_cas'] = str(cas_values['9'])
                            defaults[f'total_closed_{current_year}_joint'] = str(joint_values.get('9', '3'))
                            defaults[f'total_closed_{current_year}_subtotal'] = str(subtotal_values.get('9', '143'))
                            defaults[f'total_closed_{current_year}_court'] = str(court_values.get('9', '16'))
                            defaults[f'total_closed_{current_year}_total'] = str(total_values.get('9', '159'))
                        
                        # Row 10: TOTAL cumulatively Closed
                        if '10' in cas_values:
                            defaults['total_cumulative_closed_cas'] = str(cas_values['10'])
                            defaults['total_cumulative_closed_joint'] = str(joint_values.get('10', '16'))
                            defaults['total_cumulative_closed_subtotal'] = str(subtotal_values.get('10', '979'))
                            defaults['total_cumulative_closed_court'] = str(court_values.get('10', '77'))
                            defaults['total_cumulative_closed_total'] = str(total_values.get('10', '1056'))
                        
                        # Row 11: Total Audited (open & closed)
                        if '11' in cas_values:
                            defaults['total_audited_cas'] = str(cas_values['11'])
                            defaults['total_audited_joint'] = str(joint_values.get('11', '16'))
                            defaults['total_audited_subtotal'] = str(subtotal_values.get('11', '1122'))
                            defaults['total_audited_court'] = str(court_values.get('11', '82'))
                            defaults['total_audited_total'] = str(total_values.get('11', '1204'))
                
                # Check if we have error_rates data  
                if 'error_rates' in existing_data:
                    st.success("✅ Found existing Error Rates data - populating form!")
                    error_data = existing_data['error_rates']
                    
                    # Extract error rates values
                    if isinstance(error_data, dict):
                        error_rates = error_data.get('Error Rates (all cumulative)', {})
                        comments = error_data.get('Comments', {})
                        reporting = error_data.get('To be reported', {})
                        
                        # Map back to form fields
                        if '0' in error_rates:
                            defaults['cas_error_rate'] = str(error_rates['0']).replace('%', '')
                            defaults['cas_comments'] = str(comments.get('0', defaults['cas_comments']))
                            defaults['cas_to_be_reported'] = str(reporting.get('0', defaults['cas_to_be_reported']))
                        
                        if '1' in error_rates:
                            defaults['ercea_residual_error_rate'] = str(error_rates['1']).replace('%', '')
                            defaults['ercea_residual_comments'] = str(comments.get('1', defaults['ercea_residual_comments']))
                            defaults['ercea_residual_to_be_reported'] = str(reporting.get('1', defaults['ercea_residual_to_be_reported']))
                        
                        if '2' in error_rates:
                            defaults['ercea_overall_error_rate'] = str(error_rates['2']).replace('%', '')
                            defaults['ercea_overall_comments'] = str(comments.get('2', defaults['ercea_overall_comments']))
                            defaults['ercea_overall_to_be_reported'] = str(reporting.get('2', defaults['ercea_overall_to_be_reported']))
            else:
                st.info("📭 No existing data found, using defaults.")
                
        except Exception as e:
            st.info(f"📭 No existing data found, using defaults.")
        
        return defaults

    def collect_form_data_from_session_state():
        """Collect all form data from Streamlit session state."""
        form_data = {}
        
        # Complete list of ALL form fields
        form_fields = [
            'target_cas', 'target_joint', 'target_subtotal', 'target_court', 'target_total',
            'cumulative_cas', 'cumulative_joint', 'cumulative_subtotal', 'cumulative_court', 'cumulative_total',
            'planned_cas', 'planned_joint', 'planned_subtotal', 'planned_court', 'planned_total',
            f'ongoing_{current_year}_cas', f'ongoing_{current_year}_joint', f'ongoing_{current_year}_subtotal',
            f'ongoing_{current_year}_court', f'ongoing_{current_year}_total',
            'ongoing_prev_cas', 'ongoing_prev_joint', 'ongoing_prev_subtotal', 'ongoing_prev_court', 'ongoing_prev_total',
            f'total_ongoing_{current_year}_cas', f'total_ongoing_{current_year}_joint',
            f'total_ongoing_{current_year}_subtotal', f'total_ongoing_{current_year}_court', f'total_ongoing_{current_year}_total',
            'closed_prev_cas', 'closed_prev_joint', 'closed_prev_subtotal', 'closed_prev_court', 'closed_prev_total',
            f'audited_{current_year}_cas', f'audited_{current_year}_joint', f'audited_{current_year}_subtotal',
            f'audited_{current_year}_court', f'audited_{current_year}_total',
            f'closed_{current_year}_prev_cas', f'closed_{current_year}_prev_joint',
            f'closed_{current_year}_prev_subtotal', f'closed_{current_year}_prev_court', f'closed_{current_year}_prev_total',
            f'total_closed_{current_year}_cas', f'total_closed_{current_year}_joint',
            f'total_closed_{current_year}_subtotal', f'total_closed_{current_year}_court', f'total_closed_{current_year}_total',
            'total_cumulative_closed_cas', 'total_cumulative_closed_joint',
            'total_cumulative_closed_subtotal', 'total_cumulative_closed_court', 'total_cumulative_closed_total',
            'total_audited_cas', 'total_audited_joint', 'total_audited_subtotal',
            'total_audited_court', 'total_audited_total',
            # Error Rates fields
            'cas_error_rate', 'cas_comments', 'cas_to_be_reported',
            'ercea_residual_error_rate', 'ercea_residual_comments', 'ercea_residual_to_be_reported',
            'ercea_overall_error_rate', 'ercea_overall_comments', 'ercea_overall_to_be_reported'
        ]
        
        for field in form_fields:
            value = st.session_state.get(field, "")
            form_data[field] = value if value else "N/A"
        
        return form_data
    
    # Load existing data or set defaults
    field_values = load_existing_data_or_defaults(chosen_report, current_year)

    # Form for External Audits Table (Table 11a)
    st.subheader("External Audits (Table 11a)")
    with st.form(key="external_audits_form"):
        st.write("Enter data for each category. Use 'N/A' or leave blank for inapplicable fields.")

        # Column headers
        cols_header = st.columns([3, 1, 1.2, 2, 1.2, 1])
        headers = ["Status", "CAS", "Joint with Court*", "Subtotal for error rates", "Court only", "Total"]
        for i, header in enumerate(headers):
            with cols_header[i]:
                st.markdown(f'<div class="column-header">{header}</div>', unsafe_allow_html=True)

        # 1. ERCEA TARGETS
        st.markdown(f'<div class="section-label">{current_year} ERCEA TARGETS (Audited Participations foreseen acc. to H2020 audit strategy)</div>', unsafe_allow_html=True)
        cols1 = st.columns([3, 1, 1, 2, 1, 1])
        with cols1[0]:
            st.write("")  # Empty for alignment
        with cols1[1]:
            target_cas = st.text_input("CAS", value=field_values.get('target_cas', "150"), key="target_cas", label_visibility="collapsed")
        with cols1[2]:
            target_joint = st.text_input("Joint", value=field_values.get('target_joint', "N/A"), key="target_joint", label_visibility="collapsed")
        with cols1[3]:
            target_subtotal = st.text_input("Subtotal", value=field_values.get('target_subtotal', "150"), key="target_subtotal", label_visibility="collapsed")
        with cols1[4]:
            target_court = st.text_input("Court", value=field_values.get('target_court', "N/A"), key="target_court", label_visibility="collapsed")
        with cols1[5]:
            target_total = st.text_input("Total", value=field_values.get('target_total', "150"), key="target_total", label_visibility="collapsed")

        # 2. ERCEA TARGETS CUMULATIVE
        st.markdown('<div class="section-label">ERCEA TARGETS Cumulative</div>', unsafe_allow_html=True)
        cols2 = st.columns([3, 1, 1, 2, 1, 1])
        with cols2[0]:
            st.write("")
        with cols2[1]:
            cumulative_cas = st.text_input("CAS", value=field_values.get('cumulative_cas', "1020 (900) ***"), key="cumulative_cas", label_visibility="collapsed")
        with cols2[2]:
            cumulative_joint = st.text_input("Joint", value=field_values.get('cumulative_joint', "N/A"), key="cumulative_joint", label_visibility="collapsed")
        with cols2[3]:
            cumulative_subtotal = st.text_input("Subtotal", value=field_values.get('cumulative_subtotal', "1020 (900) ***"), key="cumulative_subtotal", label_visibility="collapsed")
        with cols2[4]:
            cumulative_court = st.text_input("Court", value=field_values.get('cumulative_court', "N/A"), key="cumulative_court", label_visibility="collapsed")
        with cols2[5]:
            cumulative_total = st.text_input("Total", value=field_values.get('cumulative_total', "1020 (900) ***"), key="cumulative_total", label_visibility="collapsed")

        # 3. PLANNED
        st.markdown('<div class="section-label">Planned</div>', unsafe_allow_html=True)
        cols3 = st.columns([3, 1, 1, 2, 1, 1])
        with cols3[0]:
            st.write("")
        with cols3[1]:
            planned_cas = st.text_input("CAS", value=field_values.get('planned_cas', "150"), key="planned_cas", label_visibility="collapsed")
        with cols3[2]:
            planned_joint = st.text_input("Joint", value=field_values.get('planned_joint', "0"), key="planned_joint", label_visibility="collapsed")
        with cols3[3]:
            planned_subtotal = st.text_input("Subtotal", value=field_values.get('planned_subtotal', "150"), key="planned_subtotal", label_visibility="collapsed")
        with cols3[4]:
            planned_court = st.text_input("Court", value=field_values.get('planned_court', "0"), key="planned_court", label_visibility="collapsed")
        with cols3[5]:
            planned_total = st.text_input("Total", value=field_values.get('planned_total', "150"), key="planned_total", label_visibility="collapsed")

        # 4. ON-GOING [LAUNCHED IN CURRENT YEAR]
        st.markdown(f'<div class="section-label">On-going [Launched in {current_year}]</div>', unsafe_allow_html=True)
        cols4 = st.columns([3, 1, 1, 2, 1, 1])
        with cols4[0]:
            st.write("")
        with cols4[1]:
            ongoing_current_cas = st.text_input("CAS", value=field_values.get(f'ongoing_{current_year}_cas', "70"), key=f"ongoing_{current_year}_cas", label_visibility="collapsed")
        with cols4[2]:
            ongoing_current_joint = st.text_input("Joint", value=field_values.get(f'ongoing_{current_year}_joint', "0"), key=f"ongoing_{current_year}_joint", label_visibility="collapsed")
        with cols4[3]:
            ongoing_current_subtotal = st.text_input("Subtotal", value=field_values.get(f'ongoing_{current_year}_subtotal', "70"), key=f"ongoing_{current_year}_subtotal", label_visibility="collapsed")
        with cols4[4]:
            ongoing_current_court = st.text_input("Court", value=field_values.get(f'ongoing_{current_year}_court', "5"), key=f"ongoing_{current_year}_court", label_visibility="collapsed")
        with cols4[5]:
            ongoing_current_total = st.text_input("Total", value=field_values.get(f'ongoing_{current_year}_total', "75"), key=f"ongoing_{current_year}_total", label_visibility="collapsed")

        # 5. ON-GOING [LAUNCHED IN PREVIOUS YEARS]
        st.markdown('<div class="section-label">On-going [Launched in previous years]</div>', unsafe_allow_html=True)
        cols5 = st.columns([3, 1, 1, 2, 1, 1])
        with cols5[0]:
            st.write("")
        with cols5[1]:
            ongoing_prev_cas = st.text_input("CAS", value=field_values.get('ongoing_prev_cas', "73"), key="ongoing_prev_cas", label_visibility="collapsed")
        with cols5[2]:
            ongoing_prev_joint = st.text_input("Joint", value=field_values.get('ongoing_prev_joint', "0"), key="ongoing_prev_joint", label_visibility="collapsed")
        with cols5[3]:
            ongoing_prev_subtotal = st.text_input("Subtotal", value=field_values.get('ongoing_prev_subtotal', "73"), key="ongoing_prev_subtotal", label_visibility="collapsed")
        with cols5[4]:
            ongoing_prev_court = st.text_input("Court", value=field_values.get('ongoing_prev_court', "0"), key="ongoing_prev_court", label_visibility="collapsed")
        with cols5[5]:
            ongoing_prev_total = st.text_input("Total", value=field_values.get('ongoing_prev_total', "73"), key="ongoing_prev_total", label_visibility="collapsed")

        # 6. TOTAL ON-GOING
        st.markdown(f'<div class="section-label">TOTAL On-going as of 31 December {current_year}</div>', unsafe_allow_html=True)
        cols6 = st.columns([3, 1, 1, 2, 1, 1])
        with cols6[0]:
            st.write("")
        with cols6[1]:
            total_ongoing_cas = st.text_input("CAS", value=field_values.get(f'total_ongoing_{current_year}_cas', "143"), key=f"total_ongoing_{current_year}_cas", label_visibility="collapsed")
        with cols6[2]:
            total_ongoing_joint = st.text_input("Joint", value=field_values.get(f'total_ongoing_{current_year}_joint', "0"), key=f"total_ongoing_{current_year}_joint", label_visibility="collapsed")
        with cols6[3]:
            total_ongoing_subtotal = st.text_input("Subtotal", value=field_values.get(f'total_ongoing_{current_year}_subtotal', "143"), key=f"total_ongoing_{current_year}_subtotal", label_visibility="collapsed")
        with cols6[4]:
            total_ongoing_court = st.text_input("Court", value=field_values.get(f'total_ongoing_{current_year}_court', "5"), key=f"total_ongoing_{current_year}_court", label_visibility="collapsed")
        with cols6[5]:
            total_ongoing_total = st.text_input("Total", value=field_values.get(f'total_ongoing_{current_year}_total', "148"), key=f"total_ongoing_{current_year}_total", label_visibility="collapsed")

        # 7. CLOSED IN PREVIOUS YEARS
        st.markdown('<div class="section-label">Closed in previous years</div>', unsafe_allow_html=True)
        cols7 = st.columns([3, 1, 1, 2, 1, 1])
        with cols7[0]:
            st.write("")
        with cols7[1]:
            closed_prev_cas = st.text_input("CAS", value=field_values.get('closed_prev_cas', "823"), key="closed_prev_cas", label_visibility="collapsed")
        with cols7[2]:
            closed_prev_joint = st.text_input("Joint", value=field_values.get('closed_prev_joint', "13"), key="closed_prev_joint", label_visibility="collapsed")
        with cols7[3]:
            closed_prev_subtotal = st.text_input("Subtotal", value=field_values.get('closed_prev_subtotal', "836"), key="closed_prev_subtotal", label_visibility="collapsed")
        with cols7[4]:
            closed_prev_court = st.text_input("Court", value=field_values.get('closed_prev_court', "61"), key="closed_prev_court", label_visibility="collapsed")
        with cols7[5]:
            closed_prev_total = st.text_input("Total", value=field_values.get('closed_prev_total', "897"), key="closed_prev_total", label_visibility="collapsed")

        # 8. CLOSED IN CURRENT YEAR FROM AUDITED PARTICIPATIONS LAUNCHED IN CURRENT YEAR
        st.markdown(f'<div class="section-label">Closed in {current_year} from audited participations launched in {current_year} (Letter of Conclusion sent)</div>', unsafe_allow_html=True)
        cols8 = st.columns([3, 1, 1, 2, 1, 1])
        with cols8[0]:
            st.write("")
        with cols8[1]:
            audited_current_cas = st.text_input("CAS", value=field_values.get(f'audited_{current_year}_cas', "8"), key=f"audited_{current_year}_cas", label_visibility="collapsed")
        with cols8[2]:
            audited_current_joint = st.text_input("Joint", value=field_values.get(f'audited_{current_year}_joint', "3"), key=f"audited_{current_year}_joint", label_visibility="collapsed")
        with cols8[3]:
            audited_current_subtotal = st.text_input("Subtotal", value=field_values.get(f'audited_{current_year}_subtotal', "11"), key=f"audited_{current_year}_subtotal", label_visibility="collapsed")
        with cols8[4]:
            audited_current_court = st.text_input("Court", value=field_values.get(f'audited_{current_year}_court', "15"), key=f"audited_{current_year}_court", label_visibility="collapsed")
        with cols8[5]:
            audited_current_total = st.text_input("Total", value=field_values.get(f'audited_{current_year}_total', "26"), key=f"audited_{current_year}_total", label_visibility="collapsed")

        # 9. CLOSED IN CURRENT YEAR FROM AUDITED PARTICIPATIONS LAUNCHED IN PREVIOUS YEARS
        st.markdown(f'<div class="section-label">Closed in {current_year} from audited participations launched in previous years</div>', unsafe_allow_html=True)
        cols9 = st.columns([3, 1, 1, 2, 1, 1])
        with cols9[0]:
            st.write("")
        with cols9[1]:
            closed_current_prev_cas = st.text_input("CAS", value=field_values.get(f'closed_{current_year}_prev_cas', "132"), key=f"closed_{current_year}_prev_cas", label_visibility="collapsed")
        with cols9[2]:
            closed_current_prev_joint = st.text_input("Joint", value=field_values.get(f'closed_{current_year}_prev_joint', "0"), key=f"closed_{current_year}_prev_joint", label_visibility="collapsed")
        with cols9[3]:
            closed_current_prev_subtotal = st.text_input("Subtotal", value=field_values.get(f'closed_{current_year}_prev_subtotal', "132"), key=f"closed_{current_year}_prev_subtotal", label_visibility="collapsed")
        with cols9[4]:
            closed_current_prev_court = st.text_input("Court", value=field_values.get(f'closed_{current_year}_prev_court', "1"), key=f"closed_{current_year}_prev_court", label_visibility="collapsed")
        with cols9[5]:
            closed_current_prev_total = st.text_input("Total", value=field_values.get(f'closed_{current_year}_prev_total', "133"), key=f"closed_{current_year}_prev_total", label_visibility="collapsed")

        # 10. TOTAL CLOSED IN CURRENT YEAR
        st.markdown(f'<div class="section-label">TOTAL Closed in {current_year}</div>', unsafe_allow_html=True)
        cols10 = st.columns([3, 1, 1, 2, 1, 1])
        with cols10[0]:
            st.write("")
        with cols10[1]:
            total_closed_cas = st.text_input("CAS", value=field_values.get(f'total_closed_{current_year}_cas', "140"), key=f"total_closed_{current_year}_cas", label_visibility="collapsed")
        with cols10[2]:
            total_closed_joint = st.text_input("Joint", value=field_values.get(f'total_closed_{current_year}_joint', "3"), key=f"total_closed_{current_year}_joint", label_visibility="collapsed")
        with cols10[3]:
            total_closed_subtotal = st.text_input("Subtotal", value=field_values.get(f'total_closed_{current_year}_subtotal', "143"), key=f"total_closed_{current_year}_subtotal", label_visibility="collapsed")
        with cols10[4]:
            total_closed_court = st.text_input("Court", value=field_values.get(f'total_closed_{current_year}_court', "16"), key=f"total_closed_{current_year}_court", label_visibility="collapsed")
        with cols10[5]:
            total_closed_total = st.text_input("Total", value=field_values.get(f'total_closed_{current_year}_total', "159"), key=f"total_closed_{current_year}_total", label_visibility="collapsed")

        # 11. TOTAL CUMULATIVELY CLOSED
        st.markdown('<div class="section-label">TOTAL cumulatively Closed</div>', unsafe_allow_html=True)
        cols11 = st.columns([3, 1, 1, 2, 1, 1])
        with cols11[0]:
            st.write("")
        with cols11[1]:
            total_cumulative_cas = st.text_input("CAS", value=field_values.get('total_cumulative_closed_cas', "963"), key="total_cumulative_closed_cas", label_visibility="collapsed")
        with cols11[2]:
            total_cumulative_joint = st.text_input("Joint", value=field_values.get('total_cumulative_closed_joint', "16"), key="total_cumulative_closed_joint", label_visibility="collapsed")
        with cols11[3]:
            total_cumulative_subtotal = st.text_input("Subtotal", value=field_values.get('total_cumulative_closed_subtotal', "979"), key="total_cumulative_closed_subtotal", label_visibility="collapsed")
        with cols11[4]:
            total_cumulative_court = st.text_input("Court", value=field_values.get('total_cumulative_closed_court', "77"), key="total_cumulative_closed_court", label_visibility="collapsed")
        with cols11[5]:
            total_cumulative_total = st.text_input("Total", value=field_values.get('total_cumulative_closed_total', "1056"), key="total_cumulative_closed_total", label_visibility="collapsed")

        # 12. TOTAL AUDITED (OPEN & CLOSED)
        st.markdown('<div class="section-label">Total Audited (open & closed) **</div>', unsafe_allow_html=True)
        cols12 = st.columns([3, 1, 1, 2, 1, 1])
        with cols12[0]:
            st.write("")
        with cols12[1]:
            total_audited_cas = st.text_input("CAS", value=field_values.get('total_audited_cas', "1106"), key="total_audited_cas", label_visibility="collapsed")
        with cols12[2]:
            total_audited_joint = st.text_input("Joint", value=field_values.get('total_audited_joint', "16"), key="total_audited_joint", label_visibility="collapsed")
        with cols12[3]:
            total_audited_subtotal = st.text_input("Subtotal", value=field_values.get('total_audited_subtotal', "1122"), key="total_audited_subtotal", label_visibility="collapsed")
        with cols12[4]:
            total_audited_court = st.text_input("Court", value=field_values.get('total_audited_court', "82"), key="total_audited_court", label_visibility="collapsed")
        with cols12[5]:
            total_audited_total = st.text_input("Total", value=field_values.get('total_audited_total', "1204"), key="total_audited_total", label_visibility="collapsed")

        submit_external_audits = st.form_submit_button("Submit External Audits Data", type="primary")

    # Form for Error Rates Table (Table 11b)
    st.subheader("Error Rates (Table 11b)")
    with st.form(key="error_rates_form"):
        st.write("Enter data for each category. Use numbers or text as needed.")

        # Column headers for Error Rates
        cols_header_error = st.columns([3, 1, 3, 1])
        headers_error = ["Name", "Error Rate (%)", "Comments", "To be reported"]
        for i, header in enumerate(headers_error):
            with cols_header_error[i]:
                st.markdown(f'<div class="column-header">{header}</div>', unsafe_allow_html=True)

        # 1. CAS CRS 1 to 6 - Latest figures
        st.markdown('<div class="section-label">CAS CRS 1 to 6 - Latest figures</div>', unsafe_allow_html=True)
        cols_error1 = st.columns([3, 1, 3, 1])
        with cols_error1[0]:
            st.write("")  # Empty for alignment
        with cols_error1[1]:
            cas_error_rate = st.text_input("Rate", value=field_values.get('cas_error_rate', "3.55"), key="cas_error_rate", label_visibility="collapsed")
        with cols_error1[2]:
            cas_comments = st.text_input("Comments", value=field_values.get('cas_comments', "Common Representative Error rate computed by the Common Audit Service (CAS) with top ups included. (source: SAR-Wiki)"), key="cas_comments", label_visibility="collapsed")
        with cols_error1[3]:
            cas_to_be_reported = st.text_input("Reporting", value=field_values.get('cas_to_be_reported', "Quarterly basis"), key="cas_to_be_reported", label_visibility="collapsed")

        # 2. ERCEA Residual Based on CRS 1 to 5 - Latest figures
        st.markdown('<div class="section-label">ERCEA Residual Based on CRS 1 to 5 - Latest figures</div>', unsafe_allow_html=True)
        cols_error2 = st.columns([3, 1, 3, 1])
        with cols_error2[0]:
            st.write("")
        with cols_error2[1]:
            ercea_residual_error_rate = st.text_input("Rate", value=field_values.get('ercea_residual_error_rate', "0.92"), key="ercea_residual_error_rate", label_visibility="collapsed")
        with cols_error2[2]:
            ercea_residual_comments = st.text_input("Comments", value=field_values.get('ercea_residual_comments', "ERCEA Residual error rate based on the CRS 1, 2, 3 & 4 (source: SAR-Wiki)"), key="ercea_residual_comments", label_visibility="collapsed")
        with cols_error2[3]:
            ercea_residual_to_be_reported = st.text_input("Reporting", value=field_values.get('ercea_residual_to_be_reported', "Quarterly basis"), key="ercea_residual_to_be_reported", label_visibility="collapsed")

        # 3. ERCEA overall detected average error rate - Latest figures
        st.markdown('<div class="section-label">ERCEA overall detected average error rate - Latest figures</div>', unsafe_allow_html=True)
        cols_error3 = st.columns([3, 1, 3, 1])
        with cols_error3[0]:
            st.write("")
        with cols_error3[1]:
            ercea_overall_error_rate = st.text_input("Rate", value=field_values.get('ercea_overall_error_rate', "1.30"), key="ercea_overall_error_rate", label_visibility="collapsed")
        with cols_error3[2]:
            ercea_overall_comments = st.text_input("Comments", value=field_values.get('ercea_overall_comments', "All ERCEA participations audited (source: SAR-Wiki)"), key="ercea_overall_comments", label_visibility="collapsed")
        with cols_error3[3]:
            ercea_overall_to_be_reported = st.text_input("Reporting", value=field_values.get('ercea_overall_to_be_reported', "Quarterly basis"), key="ercea_overall_to_be_reported", label_visibility="collapsed")

        submit_error_rates = st.form_submit_button("Submit Error Rates Data", type="primary")

    # Process External Audits Form Submission
    if submit_external_audits:
        try:
            # Collect form data
            form_data = collect_form_data_from_session_state()
            
            # Build DataFrame for External Audits
            external_audits_df = create_external_audits_dataframe(current_year, last_date, form_data)

            # Create GT Table with error handling
            try:
                external_audits_gt = GT(external_audits_df)
                external_audits_gt = apply_external_audits_styling(external_audits_gt, report_params)
            except Exception as gt_error:
                st.warning(f"⚠️ GT Table styling failed: {gt_error}")
                external_audits_gt = GT(external_audits_df)  # Use basic GT table

            # Save to Database
            insert_variable(
                report=chosen_report,
                module="AuditDataInput",
                var="external_audits",
                value=external_audits_df.to_dict(),
                db_path=DB_PATH,
                anchor="external_audits",
                gt_table=external_audits_gt
            )

            st.success("✅ External Audits data saved successfully!")
      
            # 🔍 IMMEDIATE IMAGE VERIFICATION FROM DATABASE
            st.write("### 🔍 Immediate Image Verification from Database")
            try:
                # Fetch the image that was just saved
                gt_image, anchor_name = fetch_gt_image(chosen_report, "external_audits", DB_PATH)
                
                if gt_image:
                    st.write(f"📊 **Image found in database!**")
                    st.write(f"🔍 **Image type:** {type(gt_image)}")
                    st.write(f"🔍 **Image length:** {len(str(gt_image))} characters/bytes")
                    
                    try:
                        if isinstance(gt_image, str):  # If gt_image is a file path
                            st.write(f"📁 **Image path:** {gt_image}")
                            image_path = Path(gt_image)
                            if image_path.exists():
                                st.write(f"✅ **File exists:** {image_path.stat().st_size} bytes")
                                image = Image.open(image_path)
                                st.image(image, caption=f"Saved image for {anchor_name}", use_container_width=True)
                            else:
                                st.error(f"❌ **Image file not found at:** {gt_image}")
                        else:  # If gt_image is bytes
                            st.write(f"💾 **Image is stored as bytes**")
                            image = Image.open(io.BytesIO(gt_image))
                            st.image(image, caption=f"Saved image for {anchor_name}", use_container_width=True)
                            
                        st.write(f"📍 **Anchor name:** `{anchor_name or '–'}`")
                        
                    except Exception as img_error:
                        st.error(f"❌ **Failed to display saved image:** {img_error}")
                        st.code(f"Error details: {img_error}")
                else:
                    st.error("❌ **No image found in database after save!**")
                    
            except Exception as verification_error:
                st.error(f"❌ **Image verification failed:** {verification_error}")
                import traceback
                st.code(traceback.format_exc())
            
            # Display the table in Streamlit (for comparison)
            st.write("### 📋 Streamlit GT Table Display (for comparison)")
            try:
                st.write(external_audits_gt.as_raw_html(), unsafe_allow_html=True)
            except Exception as display_error:
                st.warning(f"⚠️ GT display failed: {display_error}")
            st.dataframe(external_audits_df, use_container_width=True)


        except Exception as e:
            st.error(f"❌ Failed to process External Audits data: {str(e)}")
            import traceback
            st.code(traceback.format_exc())

    # Process Error Rates Form Submission
    if submit_error_rates:
        try:
            # Collect form data
            form_data = collect_form_data_from_session_state()
            
            # Build DataFrame for Error Rates
            error_rates_df = create_error_rates_dataframe(form_data)

            # Create GT Table with error handling
            try:
                error_rates_gt = GT(error_rates_df)
                error_rates_gt = apply_error_rates_styling(error_rates_gt, report_params)
            except Exception as gt_error:
                st.warning(f"⚠️ GT Table styling failed: {gt_error}")
                error_rates_gt = GT(error_rates_df)  # Use basic GT table

            # Save to Database
            insert_variable(
                report=chosen_report,
                module="AuditDataInput",
                var="error_rates",
                value=error_rates_df.to_dict(),
                db_path=DB_PATH,
                anchor="error_rates",
                gt_table=error_rates_gt
            )

            st.success("✅ Error Rates data saved successfully!")
            
            # Display the table
            # 🔍 IMMEDIATE IMAGE VERIFICATION FROM DATABASE
            st.write("### 🔍 Immediate Image Verification from Database")
            try:
                # Fetch the image that was just saved
                gt_image, anchor_name = fetch_gt_image(chosen_report, "error_rates", DB_PATH)
                
                if gt_image:
                    st.write(f"📊 **Image found in database!**")
                    st.write(f"🔍 **Image type:** {type(gt_image)}")
                    st.write(f"🔍 **Image length:** {len(str(gt_image))} characters/bytes")
                    
                    try:
                        if isinstance(gt_image, str):  # If gt_image is a file path
                            st.write(f"📁 **Image path:** {gt_image}")
                            image_path = Path(gt_image)
                            if image_path.exists():
                                st.write(f"✅ **File exists:** {image_path.stat().st_size} bytes")
                                image = Image.open(image_path)
                                st.image(image, caption=f"Saved image for {anchor_name}", use_container_width=True)
                            else:
                                st.error(f"❌ **Image file not found at:** {gt_image}")
                        else:  # If gt_image is bytes
                            st.write(f"💾 **Image is stored as bytes**")
                            image = Image.open(io.BytesIO(gt_image))
                            st.image(image, caption=f"Saved image for {anchor_name}", use_container_width=True)
                            
                        st.write(f"📍 **Anchor name:** `{anchor_name or '–'}`")
                        
                    except Exception as img_error:
                        st.error(f"❌ **Failed to display saved image:** {img_error}")
                        st.code(f"Error details: {img_error}")
                else:
                    st.error("❌ **No image found in database after save!**")
                    
            except Exception as verification_error:
                st.error(f"❌ **Image verification failed:** {verification_error}")
                import traceback
                st.code(traceback.format_exc())
            
            # Display the table in Streamlit (for comparison)
            st.write("### 📋 Streamlit GT Table Display (for comparison)")
            try:
                st.write(error_rates_gt.as_raw_html(), unsafe_allow_html=True)
            except Exception as display_error:
                st.warning(f"⚠️ GT display failed: {display_error}")
            st.dataframe(error_rates_df , use_container_width=True)


        except Exception as e:
            st.error(f"❌ Failed to process Error Rates data: {str(e)}")
            import traceback
            st.code(traceback.format_exc())